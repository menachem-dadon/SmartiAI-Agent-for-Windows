"""Structured Word/DOCX creation, editing, export, and visual-render tools."""
from .shared import *


_WORD_FORMATS = {
    "doc": 0,
    "dot": 1,
    "txt": 2,
    "rtf": 6,
    "html": 10,
    "htm": 10,
    "mhtml": 9,
    "docx": 16,
    "docm": 13,
    "dotx": 14,
    "dotm": 15,
    "odt": 23,
}
_WORD_FIXED_FORMATS = {"pdf": 17, "xps": 18}
_WORD_CONTENT_CONTROL_TYPES = {
    "rich_text": 0,
    "richtext": 0,
    "text": 1,
    "plain_text": 1,
    "plaintext": 1,
    "picture": 2,
    "combo_box": 3,
    "combobox": 3,
    "dropdown_list": 4,
    "dropdown": 4,
    "building_block_gallery": 5,
    "date": 6,
    "group": 7,
    "checkbox": 8,
    "check_box": 8,
    "repeating_section": 9,
}
_WORD_CHART_TYPES = {
    "area": 1,
    "area_stacked": 76,
    "area_stacked_100": 77,
    "bar_clustered": 57,
    "bar_stacked": 58,
    "bar_stacked_100": 59,
    "bubble": 15,
    "column_clustered": 51,
    "column_stacked": 52,
    "column_stacked_100": 53,
    "doughnut": -4120,
    "doughnut_exploded": 80,
    "line": 4,
    "line_markers": 65,
    "line_stacked": 63,
    "line_stacked_100": 64,
    "pie": 5,
    "pie_exploded": 69,
    "radar": -4151,
    "radar_filled": 82,
    "radar_markers": 81,
    "scatter": -4169,
    "scatter_lines": 74,
    "scatter_lines_no_markers": 75,
    "scatter_smooth": 72,
    "scatter_smooth_no_markers": 73,
}
_WORD_LEGEND_POSITIONS = {
    "bottom": -4107,
    "corner": 2,
    "left": -4131,
    "right": -4152,
    "top": -4160,
}
_SAFE_WORD_FIELD_TYPES = {
    "AUTHOR", "COMMENTS", "CREATEDATE", "DATE", "FILENAME", "FILESIZE",
    "KEYWORDS", "LASTSAVEDBY", "NUMCHARS", "NUMPAGES", "NUMWORDS", "PAGE",
    "PAGEREF", "REF", "REVNUM", "SAVEDATE", "SECTION", "SECTIONPAGES",
    "SEQ", "STYLEREF", "SUBJECT", "TIME", "TITLE", "TOC",
}
_WORD_PRINTER_LOCK = threading.Lock()
_WORD_FIND_STRING_LIMIT = 255
_WORD_BUILTIN_STYLES = {
    "normal": -1,
    "heading1": -2, "heading2": -3, "heading3": -4,
    "heading4": -5, "heading5": -6, "heading6": -7,
    "heading7": -8, "heading8": -9, "heading9": -10,
    "title": -63, "subtitle": -75, "quote": -181,
    "intensequote": -182, "listparagraph": -180,
}
_PYTHON_BLOCK_TYPES = {
    "paragraph", "heading", "title", "subtitle", "quote", "callout",
    "list", "table", "image", "page_break", "section_break", "toc", "field",
    "hyperlink", "bookmark", "header", "footer", "content_control",
}
_COM_ONLY_BLOCK_TYPES = {
    "comment", "footnote", "endnote", "text_box", "shape", "chart",
    "equation", "advanced_com",
}
_COM_ONLY_OPERATIONS = {
    "format", "format_range", "delete", "delete_range",
    "add_comment", "add_footnote", "add_endnote", "add_text_box",
    "add_shape", "add_chart", "add_equation", "track_changes",
    "accept_all_changes", "reject_all_changes", "protect", "unprotect",
    "advanced_com", "insert_file",
}
_ADVANCED_COM_BLOCKED = {
    "vbproject", "vbe", "macro", "macros", "run", "ddeexecute",
    "ddeinitiate", "dderequest", "oleobjects", "addoleobject", "addolecontrol",
    "commandbars", "addins", "comaddins", "followhyperlink", "sendmail",
    "route", "routingslip", "printout", "printpreview", "saveas", "saveas2",
    "exportasfixedformat", "exportasfixedformat2", "open", "opentext",
    "openold", "attachedtemplate", "linkformat", "executeexcel4macro",
    "save", "close", "quit", "wordbasic", "system", "tasks", "filedialog",
    "organizercopy", "organizerdelete", "sendforreview", "sendfaxoverinternet",
    "reply", "replyall", "forward",
}


def _json_text(payload):
    return json.dumps(payload, ensure_ascii=False, indent=2, default=str)


def _clamp(value, minimum, maximum, default):
    try:
        return max(minimum, min(maximum, float(value)))
    except Exception:
        return default


def _literal_match_spans(text, needle, *, case_sensitive=False, whole_word=False):
    """Return literal match offsets without relying on Word Find limits."""
    text = str(text or "")
    needle = str(needle or "")
    if not needle:
        return []
    pattern = re.escape(needle)
    if whole_word:
        pattern = rf"(?<!\w){pattern}(?!\w)"
    flags = 0 if case_sensitive else re.IGNORECASE
    return [match.span() for match in re.finditer(pattern, text, flags)]


def _merge_nested_spec(spec, *keys):
    """Flatten common model-friendly wrappers without making them engine-specific."""
    merged = dict(spec) if isinstance(spec, dict) else {}
    for key in keys:
        nested = merged.get(key)
        if isinstance(nested, dict):
            merged.update(nested)
    return merged


def _normalized_format_spec(spec):
    merged = _merge_nested_spec(spec, "format")
    aliases = {
        "font_name": "font",
        "font_size": "font_size_pt",
        "style_type": "type",
        "space_before": "space_before_pt",
        "space_after": "space_after_pt",
        "left_indent": "left_indent_cm",
        "right_indent": "right_indent_cm",
        "first_line_indent": "first_line_indent_cm",
    }
    for alias, canonical in aliases.items():
        if canonical not in merged and alias in merged:
            merged[canonical] = merged[alias]
    return merged


def _safe_color(value, default="000000"):
    text = str(value or default).strip().lstrip("#")
    return text.upper() if re.fullmatch(r"[0-9a-fA-F]{6}", text) else default


def _safe_com_result(value, depth=0):
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    if depth >= 2:
        return str(value)[:500]
    if isinstance(value, (list, tuple)):
        return [_safe_com_result(item, depth + 1) for item in value[:100]]
    if isinstance(value, dict):
        return {
            str(key)[:100]: _safe_com_result(item, depth + 1)
            for key, item in list(value.items())[:100]
        }
    for attr in ("Name", "Title", "Text", "Count"):
        try:
            item = getattr(value, attr)
            if not callable(item):
                return {"type": type(value).__name__, attr.lower(): str(item)[:1000]}
        except Exception:
            pass
    return {"type": type(value).__name__, "value": str(value)[:1000]}


def _friendly_enum(value, values, label, default):
    if value is None or value == "":
        return default
    if isinstance(value, bool):
        raise ValueError(f"{label} must be a friendly name or numeric Office constant.")
    if isinstance(value, (int, float)) and int(value) == value:
        return int(value)
    text = re.sub(r"[^a-z0-9]+", "_", str(value).strip().lower()).strip("_")
    if re.fullmatch(r"-?\d+", text):
        return int(text)
    if text in values:
        return values[text]
    allowed = ", ".join(sorted(values))
    raise ValueError(f"Unsupported {label} '{value}'. Use a numeric Office constant or one of: {allowed}.")


def _bookmark_name(value):
    clean = re.sub(r"[^A-Za-z0-9_]", "_", str(value or "").strip())[:40]
    return clean or "bookmark"


def _pymupdf_module():
    try:
        import pymupdf
        return pymupdf
    except Exception:
        try:
            import fitz
            return fitz
        except Exception as exc:
            raise RuntimeError(
                "PyMuPDF is required for PDF inspection and page-PNG visual QA. "
                "Install Smarti requirements (python -m pip install PyMuPDF)."
            ) from exc


def _safe_layout_printer():
    """Return a local virtual printer Word can use without contacting a WSD device."""
    if platform.system() != "Windows":
        return None
    try:
        import win32print
        flags = win32print.PRINTER_ENUM_LOCAL | win32print.PRINTER_ENUM_CONNECTIONS
        installed = {
            str(item[2]).strip().casefold(): str(item[2]).strip()
            for item in win32print.EnumPrinters(flags)
            if len(item) > 2 and str(item[2]).strip()
        }
        printer_specs = []
        for name in installed.values():
            handle = win32print.OpenPrinter(name)
            try:
                port = str(win32print.GetPrinter(handle, 2).get("pPortName") or "").strip()
            finally:
                win32print.ClosePrinter(handle)
            printer_specs.append((name, port))
        # A local null-port driver is ideal for layout: it never prompts for a
        # file name. XPS/PDF remain fallbacks on systems without OneNote.
        preferred_names = (
            "OneNote (Desktop)", "Microsoft XPS Document Writer", "Microsoft Print to PDF",
        )
        for name, port in printer_specs:
            if port.casefold() == "nul:":
                return f"{name} on {port}"
        for preferred in preferred_names:
            for name, port in printer_specs:
                if name.casefold() == preferred.casefold():
                    return f"{name} on {port}" if port else name
    except Exception:
        logging.debug("Could not enumerate local layout printers.", exc_info=True)
    return None


def _winword_pids():
    if platform.system() != "Windows":
        return set()
    result = set()
    try:
        import ctypes
        from ctypes import wintypes

        class PROCESSENTRY32W(ctypes.Structure):
            _fields_ = [
                ("dwSize", wintypes.DWORD),
                ("cntUsage", wintypes.DWORD),
                ("th32ProcessID", wintypes.DWORD),
                ("th32DefaultHeapID", wintypes.WPARAM),
                ("th32ModuleID", wintypes.DWORD),
                ("cntThreads", wintypes.DWORD),
                ("th32ParentProcessID", wintypes.DWORD),
                ("pcPriClassBase", wintypes.LONG),
                ("dwFlags", wintypes.DWORD),
                ("szExeFile", wintypes.WCHAR * 260),
            ]

        kernel32 = ctypes.WinDLL("kernel32", use_last_error=True)
        kernel32.CreateToolhelp32Snapshot.argtypes = (wintypes.DWORD, wintypes.DWORD)
        kernel32.CreateToolhelp32Snapshot.restype = wintypes.HANDLE
        kernel32.Process32FirstW.argtypes = (wintypes.HANDLE, ctypes.POINTER(PROCESSENTRY32W))
        kernel32.Process32FirstW.restype = wintypes.BOOL
        kernel32.Process32NextW.argtypes = (wintypes.HANDLE, ctypes.POINTER(PROCESSENTRY32W))
        kernel32.Process32NextW.restype = wintypes.BOOL
        kernel32.CloseHandle.argtypes = (wintypes.HANDLE,)
        kernel32.CloseHandle.restype = wintypes.BOOL
        snapshot = kernel32.CreateToolhelp32Snapshot(0x00000002, 0)  # TH32CS_SNAPPROCESS
        invalid_handle = ctypes.c_void_p(-1).value
        if snapshot == invalid_handle:
            return result
        try:
            entry = PROCESSENTRY32W()
            entry.dwSize = ctypes.sizeof(entry)
            more = kernel32.Process32FirstW(snapshot, ctypes.byref(entry))
            while more:
                if str(entry.szExeFile).casefold() == "winword.exe":
                    result.add(int(entry.th32ProcessID))
                more = kernel32.Process32NextW(snapshot, ctypes.byref(entry))
        finally:
            kernel32.CloseHandle(snapshot)
    except Exception:
        logging.debug("Could not enumerate WINWORD processes.", exc_info=True)
    return result


class _WordComSession:
    """Owns one isolated WINWORD instance and only ever kills that instance."""

    def __init__(self, timeout_seconds=180, visible=False):
        self.timeout_seconds = int(_clamp(timeout_seconds, 15, 1800, 180))
        self.visible = bool(visible)
        self.pythoncom = None
        self.app = None
        self.document = None
        self.pid = None
        self.timer = None
        self.timed_out = False
        self.layout_printer = None
        self.printer_isolated = False
        self._word_pids_before = set()

    def __enter__(self):
        try:
            import pythoncom
            import win32com.client
        except Exception as exc:
            raise RuntimeError("Word COM requires pywin32 on Windows.") from exc
        self.pythoncom = pythoncom
        pythoncom.CoInitialize()
        try:
            self._word_pids_before = _winword_pids()
            self.app = self._launch_word_with_layout_printer(win32com.client)
            for _attempt in range(10):
                new_pids = _winword_pids() - self._word_pids_before
                if len(new_pids) == 1:
                    self.pid = next(iter(new_pids))
                    break
                threading.Event().wait(0.05)
            self.app.Visible = self.visible
            self.app.DisplayAlerts = 0
            # msoAutomationSecurityForceDisable. Never execute document macros.
            self.app.AutomationSecurity = 3
            try:
                self.app.Options.ConfirmConversions = False
                self.app.Options.SaveInterval = 0
            except Exception:
                pass
            self.timer = threading.Timer(self.timeout_seconds, self._on_timeout)
            self.timer.daemon = True
            self.timer.start()
            return self
        except Exception:
            if self.app is not None:
                try:
                    self.app.Quit(SaveChanges=0)
                except Exception:
                    pass
                self.app = None
            pythoncom.CoUninitialize()
            raise

    def _launch_word_with_layout_printer(self, win32_client):
        self.layout_printer = _safe_layout_printer()
        return win32_client.DispatchEx("Word.Application")

    def _configure_layout_printer(self):
        """Select a local virtual printer in Word without changing Windows."""
        if not self.layout_printer:
            return
        try:
            import pythoncom
            import win32print
            printer_name = self.layout_printer.split(" on ", 1)[0].strip()
            with _WORD_PRINTER_LOCK:
                original_default = win32print.GetDefaultPrinter()
                dialog = self.app.Dialogs(97)  # wdDialogFilePrintSetup
                ole = dialog._oleobj_
                ole.Invoke(
                    ole.GetIDsOfNames("Printer"), 0,
                    pythoncom.DISPATCH_PROPERTYPUT, 0,
                    self.layout_printer,
                )
                ole.Invoke(
                    ole.GetIDsOfNames("DoNotSetAsSysDefault"), 0,
                    pythoncom.DISPATCH_PROPERTYPUT, 0,
                    True,
                )
                ole.Invoke(
                    ole.GetIDsOfNames("Execute"), 0,
                    pythoncom.DISPATCH_METHOD, 1,
                )
                current_default = win32print.GetDefaultPrinter()
                if current_default != original_default:
                    win32print.SetDefaultPrinter(original_default)
                    raise RuntimeError("Word changed the Windows default printer unexpectedly; it was restored.")
                active = str(self.app.ActivePrinter or "").strip()
                self.printer_isolated = active.casefold().startswith(printer_name.casefold())
                if not self.printer_isolated:
                    raise RuntimeError(f"Word kept an unexpected active printer: {active or 'unknown'}")
        except Exception:
            logging.warning(
                "Could not isolate Word layout from the default printer; continuing without printer override.",
                exc_info=True,
            )

    def _on_timeout(self):
        self.timed_out = True
        self._terminate_owned_word()

    def _terminate_owned_word(self):
        if not self.pid:
            return
        if self.pid not in _winword_pids() or self.pid in self._word_pids_before:
            return
        try:
            import win32api
            import win32con
            handle = win32api.OpenProcess(win32con.PROCESS_TERMINATE, False, int(self.pid))
            try:
                win32api.TerminateProcess(handle, 1)
            finally:
                handle.Close()
        except Exception:
            logging.exception("Could not terminate Smarti-owned WINWORD after timeout.")

    def new_document(self, template_path=""):
        kwargs = {"NewTemplate": False, "DocumentType": 0, "Visible": self.visible}
        if template_path:
            kwargs["Template"] = template_path
        self.document = self.app.Documents.Add(**kwargs)
        self._capture_document_pid()
        self._configure_layout_printer()
        return self.document

    def open_document(self, path, read_only=False, password=""):
        kwargs = {
            "FileName": path,
            "ConfirmConversions": False,
            "ReadOnly": bool(read_only),
            "AddToRecentFiles": False,
            "Visible": self.visible,
            "OpenAndRepair": True,
            "NoEncodingDialog": True,
        }
        if password:
            kwargs["PasswordDocument"] = password
        self.document = self.app.Documents.Open(**kwargs)
        self._capture_document_pid()
        self._configure_layout_printer()
        return self.document

    def _capture_document_pid(self):
        if self.pid or self.document is None:
            return
        try:
            import win32process
            hwnd = int(self.document.ActiveWindow.Hwnd)
            _thread_id, self.pid = win32process.GetWindowThreadProcessId(hwnd)
        except Exception:
            pass

    def close_document(self, save_changes=False):
        if self.document is None:
            return
        try:
            self.document.Close(SaveChanges=-1 if save_changes else 0)
        except Exception:
            pass
        self.document = None

    def __exit__(self, exc_type, exc, tb):
        if self.timer:
            self.timer.cancel()
        self.close_document(False)
        if self.app is not None:
            try:
                self.app.Quit(SaveChanges=0)
            except Exception:
                logging.warning("Could not close the Smarti-owned Word instance cleanly.", exc_info=True)
        self.app = None
        if self.pid:
            for _attempt in range(10):
                if self.pid not in _winword_pids():
                    break
                threading.Event().wait(0.05)
            self._terminate_owned_word()
        if self.pythoncom is not None:
            try:
                self.pythoncom.CoUninitialize()
            except Exception:
                pass
        if self.timed_out and exc_type is None:
            raise TimeoutError(f"Microsoft Word exceeded the {self.timeout_seconds}-second timeout.")
        return False


class DocumentToolsMixin:
    """Safe, structured document automation with COM and python-docx engines."""

    def document_manager_tool(self, args):
        args = args if isinstance(args, dict) else {}
        action = str(args.get("action") or "").strip().lower()
        if action not in {"doctor", "create", "edit", "inspect", "render", "visual_qa", "export", "compare"}:
            return "ERROR: document_manager action must be doctor, create, edit, inspect, render, visual_qa, export, or compare."
        try:
            if action == "doctor":
                return _json_text(self._document_doctor())
            if action == "visual_qa":
                path = self._document_resolve_path(args.get("path"), mode="read")
                allowed, error = self._ensure_cloud_upload_allowed(path)
                if not allowed:
                    return error
                return self._document_visual_qa(path)

            engine = self._document_choose_engine(action, args)
            self._document_validate_request(action, args, engine)
            capabilities = ["file_read"] if action in {"inspect", "render", "export", "compare", "edit"} else []
            if action in {"create", "edit", "render", "export", "compare"}:
                capabilities.append("file_write")
            if engine in {"com", "libreoffice"} or action == "compare":
                capabilities.append("office_automation")
            if action in {"create", "edit"} and (args.get("render_after") or args.get("export_formats")):
                post_engine = self._document_choose_engine(
                    "export",
                    {"path": args.get("output_path") or args.get("path"), "format": "pdf", "engine": "auto"},
                )
                if post_engine in {"com", "libreoffice"}:
                    capabilities.append("office_automation")
            advanced = self._document_has_advanced_com(args)
            risk = "high" if advanced or action == "compare" else "medium"
            details = self._document_permission_details(action, args, engine)
            allowed, error = self._ensure_capabilities_allowed(
                capabilities,
                "אישור עבודה עם מסמך Word",
                details,
                risk=risk,
                audit_context={"tool": "document_manager", "action": action, "engine": engine},
            )
            if not allowed:
                return error

            if action == "create":
                return _json_text(self._document_create(args, engine))
            if action == "edit":
                return _json_text(self._document_edit(args, engine))
            if action == "inspect":
                return _json_text(self._document_inspect(args, engine))
            if action == "render":
                return _json_text(self._document_render(args, engine))
            if action == "export":
                return _json_text(self._document_export(args, engine))
            return _json_text(self._document_compare(args))
        except Exception as exc:
            logging.exception("document_manager failed")
            return f"ERROR: document_manager {action} failed: {exc}"

    @staticmethod
    def _document_visual_qa(path):
        if os.path.splitext(path)[1].lower() != ".png":
            raise ValueError("visual_qa requires a page PNG produced by document_manager render.")
        if os.path.getsize(path) > 12 * 1024 * 1024:
            raise ValueError("Rendered page image is too large for model visual QA (max 12MB).")
        with open(path, "rb") as image_file:
            encoded = base64.b64encode(image_file.read()).decode("ascii")
        return f"IMAGE_BASE64:image/png:{encoded}"

    def _document_permission_details(self, action, args, engine):
        paths = [
            args.get("path"), args.get("output_path"), args.get("template_path"),
            args.get("other_path"), args.get("output_dir"),
        ]
        visible_paths = "\n".join(str(item) for item in paths if item)
        return f"פעולה: {action}\nמנוע: {engine}\n{visible_paths}".strip()

    def _document_doctor(self):
        word = self._word_com_available()
        libreoffice = self._find_soffice()
        try:
            import docx
            python_docx = getattr(docx, "__version__", "installed")
        except Exception:
            python_docx = None
        try:
            fitz = _pymupdf_module()
            pymupdf = getattr(fitz, "VersionBind", None) or getattr(fitz, "__version__", "installed")
        except Exception:
            pymupdf = None
        printer = self._word_printer_diagnostics()
        return {
            "status": "ok",
            "platform": platform.system(),
            "word_com": word,
            "python_docx": python_docx,
            "libreoffice": libreoffice,
            "pymupdf": pymupdf,
            "word_layout_printer": printer,
            "default_language": "he-IL",
            "default_direction": "rtl",
            "ui_automation": False,
            "security": {
                "macros": "force-disabled",
                "vba_ole_and_external_links": "blocked",
                "existing_document_backup": "enabled-by-default",
            },
        }

    @staticmethod
    def _word_printer_diagnostics():
        diagnostics = {
            "default": None,
            "isolated_to": _safe_layout_printer(),
            "changes_windows_default": False,
        }
        if platform.system() != "Windows":
            return diagnostics
        try:
            import win32print
            diagnostics["default"] = win32print.GetDefaultPrinter()
        except Exception:
            pass
        return diagnostics

    @staticmethod
    def _word_com_available():
        if platform.system() != "Windows":
            return False
        try:
            import winreg
            with winreg.OpenKey(winreg.HKEY_CLASSES_ROOT, r"Word.Application\CLSID") as key:
                value, _kind = winreg.QueryValueEx(key, None)
            return bool(str(value).strip())
        except Exception:
            return False

    @staticmethod
    def _find_soffice():
        candidates = [
            shutil.which("soffice"),
            shutil.which("libreoffice"),
            os.path.join(os.environ.get("PROGRAMFILES", ""), "LibreOffice", "program", "soffice.exe"),
            os.path.join(os.environ.get("PROGRAMFILES(X86)", ""), "LibreOffice", "program", "soffice.exe"),
        ]
        return next((os.path.abspath(item) for item in candidates if item and os.path.isfile(item)), None)

    def _document_choose_engine(self, action, args):
        requested = str(args.get("engine") or "auto").strip().lower()
        if requested not in {"auto", "com", "python", "libreoffice"}:
            raise ValueError("engine must be auto, com, python, or libreoffice.")
        if requested != "auto":
            if requested == "com" and not self._word_com_available():
                raise RuntimeError("Microsoft Word COM is unavailable. Install/licence Microsoft Word or use engine='python'.")
            if requested == "libreoffice" and not self._find_soffice():
                raise RuntimeError("LibreOffice/soffice is unavailable.")
            return requested
        if action == "compare":
            if not self._word_com_available():
                raise RuntimeError("Document comparison currently requires Microsoft Word COM.")
            return "com"
        source_ext = os.path.splitext(str(args.get("path") or ""))[1].lower()
        output_ext = os.path.splitext(str(args.get("output_path") or ""))[1].lower()
        template_ext = os.path.splitext(str(args.get("template_path") or ""))[1].lower()
        if action in {"inspect", "render"} and source_ext == ".pdf":
            return "python"
        if action == "inspect" and source_ext not in {"", ".docx"}:
            if not self._word_com_available():
                raise RuntimeError("Inspecting legacy/template Word formats requires Microsoft Word COM.")
            return "com"
        if action == "edit" and (source_ext not in {"", ".docx"} or output_ext not in {"", ".docx"}):
            if not self._word_com_available():
                raise RuntimeError("Editing legacy/template Word formats requires Microsoft Word COM.")
            return "com"
        if action == "create" and (output_ext not in {"", ".docx"} or template_ext not in {"", ".docx"}):
            if not self._word_com_available():
                raise RuntimeError("Creating legacy/template Word formats requires Microsoft Word COM.")
            return "com"
        if action in {"render", "export"}:
            fmt = str(args.get("format") or "pdf").lower().lstrip(".")
            if fmt in {"pdf", "xps", "rtf", "html", "htm", "odt", "doc"} and self._word_com_available():
                return "com"
            if self._find_soffice():
                return "libreoffice"
        if self._document_needs_com(args):
            if not self._word_com_available():
                raise RuntimeError("The requested Word feature requires Microsoft Word COM, but Word is unavailable.")
            return "com"
        return "python"

    @staticmethod
    def _document_needs_com(args):
        plan = args.get("document") if isinstance(args.get("document"), dict) else {}
        blocks = plan.get("blocks") or plan.get("sections") or []
        for block in blocks if isinstance(blocks, list) else []:
            if isinstance(block, dict) and str(block.get("type") or "paragraph").lower() in _COM_ONLY_BLOCK_TYPES:
                return True
        for operation in args.get("operations") or []:
            if isinstance(operation, dict) and str(operation.get("op") or operation.get("type") or "").lower() in _COM_ONLY_OPERATIONS:
                return True
        return False

    @staticmethod
    def _document_has_advanced_com(args):
        for operation in args.get("operations") or []:
            if isinstance(operation, dict) and str(operation.get("op") or operation.get("type") or "").lower() == "advanced_com":
                return True
        plan = args.get("document") if isinstance(args.get("document"), dict) else {}
        return any(
            isinstance(block, dict) and str(block.get("type") or "").lower() == "advanced_com"
            for block in (plan.get("blocks") or [])
        )

    def _document_validate_request(self, action, args, engine):
        if action == "create":
            plan = args.get("document") if isinstance(args.get("document"), dict) else {}
            self._document_validate_blocks(plan.get("blocks") or plan.get("sections") or [], engine)
        elif action == "edit":
            for operation in args.get("operations") or []:
                if not isinstance(operation, dict):
                    raise ValueError("edit.operations entries must be objects.")
                op = str(operation.get("op") or operation.get("type") or "").lower()
                if op in {"append", "append_blocks", "insert", "insert_blocks"}:
                    self._document_validate_blocks(operation.get("blocks") or [], engine)
                elif op in {"add_text_box", "add_shape", "add_chart", "add_equation"}:
                    self._document_validate_blocks([{**operation, "type": op.replace("add_", "")}], engine)

    def _document_validate_blocks(self, blocks, engine):
        if not isinstance(blocks, list):
            raise ValueError("document.blocks must be an array.")
        supported = _PYTHON_BLOCK_TYPES if engine == "python" else _PYTHON_BLOCK_TYPES | _COM_ONLY_BLOCK_TYPES
        for index, block in enumerate(blocks):
            if not isinstance(block, dict):
                raise ValueError(f"document.blocks[{index}] must be an object.")
            kind = str(block.get("type") or "paragraph").strip().lower()
            if kind not in supported:
                suffix = "; use engine='com'" if kind in _COM_ONLY_BLOCK_TYPES else ""
                raise ValueError(f"Unsupported block type '{kind}' for engine='{engine}'{suffix}.")
            if kind == "content_control" and engine == "com":
                _friendly_enum(block.get("control_type"), _WORD_CONTENT_CONTROL_TYPES, "control_type", 0)
            elif kind == "chart":
                _friendly_enum(block.get("chart_type"), _WORD_CHART_TYPES, "chart_type", 51)
                self._validate_chart_data(block)
            elif kind == "hyperlink":
                self._hyperlink_target(block.get("url"), block.get("anchor"))
            elif kind == "field":
                self._word_field_code(block)
            for run in block.get("runs") or []:
                if not isinstance(run, dict):
                    continue
                if run.get("hyperlink"):
                    self._hyperlink_target(run.get("hyperlink"), run.get("anchor"))
                if run.get("field"):
                    self._word_field_code({"code": run.get("field")})

    @staticmethod
    def _validate_chart_data(block):
        categories = block.get("categories") or []
        series = block.get("series") or []
        if categories and not isinstance(categories, list):
            raise ValueError("chart.categories must be an array.")
        if series and not isinstance(series, list):
            raise ValueError("chart.series must be an array.")
        if block.get("legend_position") is not None:
            _friendly_enum(block.get("legend_position"), _WORD_LEGEND_POSITIONS, "legend_position", -4107)
        for index, item in enumerate(series):
            if not isinstance(item, dict):
                raise ValueError(f"chart.series[{index}] must be an object.")
            values = item.get("values") or []
            if not isinstance(values, list):
                raise ValueError(f"chart.series[{index}].values must be an array.")
            if categories and values and len(values) != len(categories):
                raise ValueError(
                    f"chart.series[{index}] has {len(values)} values but chart.categories has {len(categories)} entries."
                )

    @staticmethod
    def _hyperlink_target(url=None, anchor=None):
        url = str(url or "").strip()
        anchor = str(anchor or "").strip()
        if url.startswith("#") and not anchor:
            anchor, url = url[1:], ""
        if anchor:
            return "", _bookmark_name(anchor)
        if re.match(r"^(https?|mailto):", url, flags=re.I):
            return url, ""
        raise ValueError("Hyperlinks require an http/https/mailto URL or an internal bookmark anchor.")

    @staticmethod
    def _word_field_code(spec):
        spec = spec if isinstance(spec, dict) else {}
        code = str(spec.get("code") or "").strip()
        field_type = str(spec.get("field_type") or "").strip().upper()
        if code:
            match = re.match(r"^\s*([A-Za-z]+)", code)
            field_type = match.group(1).upper() if match else ""
            if len(code) > 500 or any(char in code for char in ("\r", "\n", "\x00")):
                raise ValueError("Word field code is too long or contains invalid control characters.")
        if field_type not in _SAFE_WORD_FIELD_TYPES:
            allowed = ", ".join(sorted(_SAFE_WORD_FIELD_TYPES))
            raise ValueError(f"Unsupported or unsafe Word field type '{field_type or code}'. Allowed: {allowed}.")
        if code:
            return code
        parts = [field_type]
        target = str(spec.get("bookmark") or spec.get("anchor") or spec.get("target") or "").strip()
        if field_type in {"REF", "PAGEREF", "STYLEREF", "SEQ"}:
            if not target:
                raise ValueError(f"field_type={field_type} requires bookmark/anchor/target.")
            parts.append(_bookmark_name(target))
        fmt = str(spec.get("format") or "").replace('"', "").replace("\r", " ").replace("\n", " ")[:100]
        if fmt:
            switch = "\\@" if field_type in {"CREATEDATE", "DATE", "SAVEDATE", "TIME"} else "\\#"
            parts.extend((switch, f'"{fmt}"'))
        if field_type == "FILENAME" and bool(spec.get("include_path")):
            parts.append("\\p")
        return " ".join(parts)

    def _document_resolve_path(self, value, *, mode="read", default_name=""):
        text = os.path.expandvars(os.path.expanduser(str(value or "").strip().strip('"')))
        if not text:
            if not default_name:
                raise ValueError("A file path is required.")
            text = default_name
        if not os.path.isabs(text):
            base = self._sandbox_root() if self._sandbox_enabled() else OUTPUTS_DIR
            text = os.path.join(base, text)
        path = os.path.abspath(text)
        sandbox_ok, sandbox_error = self._ensure_sandbox_path_allowed(path, "read" if mode == "read" else "write")
        if not sandbox_ok:
            raise PermissionError(sandbox_error)
        if mode == "read" and not os.path.isfile(path):
            raise FileNotFoundError(path)
        if mode != "read":
            os.makedirs(os.path.dirname(path), exist_ok=True)
        return path

    @staticmethod
    def _document_temp_path(output_path):
        root, ext = os.path.splitext(output_path)
        return f"{root}.smarti-{uuid.uuid4().hex[:10]}{ext or '.docx'}"

    @staticmethod
    def _document_backup(path):
        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        root, ext = os.path.splitext(path)
        candidate = f"{root}.backup-{stamp}{ext}"
        counter = 2
        while os.path.exists(candidate):
            candidate = f"{root}.backup-{stamp}-{counter}{ext}"
            counter += 1
        shutil.copy2(path, candidate)
        return candidate

    def _document_prepare_output(self, path, overwrite=False, backup_existing=False):
        if not os.path.exists(path):
            return None
        if not overwrite:
            raise FileExistsError(f"Output already exists: {path}. Set overwrite=true to replace it.")
        return self._document_backup(path) if backup_existing else None

    def _document_create(self, args, engine):
        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        output = self._document_resolve_path(
            args.get("output_path") or args.get("path"),
            mode="write",
            default_name=f"document-{stamp}.docx",
        )
        if os.path.splitext(output)[1].lower() not in {".docx", ".docm", ".dotx", ".dotm", ".rtf", ".doc"}:
            raise ValueError("create output must be a Word document path (.docx/.docm/.dotx/.dotm/.rtf/.doc).")
        replaced_backup = self._document_prepare_output(output, bool(args.get("overwrite")), backup_existing=True)
        template = ""
        if args.get("template_path"):
            template = self._document_resolve_path(args.get("template_path"), mode="read")
        plan = args.get("document") if isinstance(args.get("document"), dict) else {}
        if engine == "com":
            result = self._com_create_document(output, template, plan, args)
        elif engine == "python":
            if os.path.splitext(output)[1].lower() != ".docx" or (template and os.path.splitext(template)[1].lower() != ".docx"):
                raise ValueError("The independent python engine creates DOCX from DOCX templates only; use engine='com' for other Word formats.")
            result = self._python_create_document(output, template, plan)
        else:
            raise ValueError("LibreOffice is an export/render engine; use python or com for create.")
        result.update({"engine": engine, "output_path": output, "replaced_file_backup": replaced_backup})
        post_actions = self._document_post_actions(output, args)
        result.update(post_actions)
        if post_actions.get("post_action_errors"):
            result["status"] = "created_with_warnings"
        return result

    def _document_edit(self, args, engine):
        source = self._document_resolve_path(args.get("path"), mode="read")
        output = self._document_resolve_path(args.get("output_path") or source, mode="write")
        in_place = os.path.normcase(source) == os.path.normcase(output)
        backup = None
        if in_place and args.get("backup", True):
            backup = self._document_backup(source)
        elif not in_place:
            self._document_prepare_output(output, bool(args.get("overwrite")), backup_existing=True)
        operations = args.get("operations") or []
        if not isinstance(operations, list) or not operations:
            raise ValueError("edit requires a non-empty operations array.")
        if engine == "com":
            result = self._com_edit_document(source, output, operations, args)
        elif engine == "python":
            if os.path.splitext(source)[1].lower() != ".docx" or os.path.splitext(output)[1].lower() != ".docx":
                raise ValueError("The independent python edit engine supports DOCX only; use engine='com' for other Word formats.")
            result = self._python_edit_document(source, output, operations)
        else:
            raise ValueError("LibreOffice is an export/render engine; use python or com for edit.")
        result.update({"engine": engine, "output_path": output, "backup_path": backup})
        post_actions = self._document_post_actions(output, args)
        result.update(post_actions)
        if post_actions.get("post_action_errors"):
            result["status"] = "edited_with_warnings"
        return result

    def _document_post_actions(self, path, args):
        payload = {}
        errors = []
        exported_pdf = None
        exports = args.get("export_formats") or []
        if isinstance(exports, str):
            exports = [exports]
        if exports:
            payload["exports"] = []
            for fmt in exports[:8]:
                fmt = str(fmt).lower().lstrip(".")
                target = os.path.splitext(path)[0] + "." + fmt
                if os.path.normcase(target) == os.path.normcase(path):
                    target = os.path.splitext(path)[0] + "-export." + fmt
                try:
                    self._document_prepare_output(target, bool(args.get("overwrite")), backup_existing=True)
                    engine = self._document_choose_engine("export", {"path": path, "format": fmt, "engine": "auto"})
                    export_info = self._document_export_internal(path, target, fmt, engine, args)
                    payload["exports"].append(export_info)
                    if fmt == "pdf":
                        exported_pdf = export_info.get("output_path")
                except Exception as exc:
                    logging.exception("Document post-create/edit export failed for format %s", fmt)
                    errors.append({"stage": "export", "format": fmt, "error": str(exc)[:1000]})
        if args.get("render_after"):
            render_args = {
                "path": exported_pdf or path,
                "output_dir": args.get("output_dir"),
                "dpi": args.get("dpi", 144),
                "include_pdf": True,
                "page_limit": args.get("page_limit", 100),
            }
            try:
                render_engine = self._document_choose_engine("render", {**render_args, "engine": "auto"})
                payload["render"] = self._document_render(render_args, render_engine)
            except Exception as exc:
                logging.exception("Document post-create/edit visual render failed")
                errors.append({"stage": "render", "error": str(exc)[:1000]})
        if errors:
            payload["post_action_errors"] = errors
            payload["warnings"] = [
                "The Word document was saved successfully, but one or more requested post-actions failed. "
                "Use the returned output_path instead of recreating the document."
            ]
        return payload

    # ------------------------- python-docx engine -------------------------

    @staticmethod
    def _python_imports():
        try:
            import docx
            from docx.enum.section import WD_ORIENT, WD_SECTION_START
            from docx.enum.style import WD_STYLE_TYPE
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Cm, Pt, RGBColor
        except Exception as exc:
            raise RuntimeError("python-docx is required for the independent DOCX engine.") from exc
        return locals()

    def _python_create_document(self, output, template, plan):
        api = self._python_imports()
        docx = api["docx"]
        document = docx.Document(template or None)
        self._python_apply_plan(document, plan, api)
        temp = self._document_temp_path(output)
        try:
            document.save(temp)
            os.replace(temp, output)
        finally:
            if os.path.exists(temp):
                os.remove(temp)
        return {"status": "created", "blocks": len(plan.get("blocks") or plan.get("sections") or [])}

    def _python_apply_plan(self, document, plan, api):
        self._python_set_core_properties(document, plan.get("metadata") or {})
        self._python_configure_defaults(document, plan, api)
        self._python_set_page_layout(document, plan.get("page") or {}, api)
        for style in plan.get("styles") or []:
            if isinstance(style, dict):
                self._python_define_style(document, style, api)
        if isinstance(plan.get("header"), dict):
            self._python_set_header_footer(document, "header", plan["header"], api)
        if isinstance(plan.get("footer"), dict):
            self._python_set_header_footer(document, "footer", plan["footer"], api)
        self._python_add_blocks(document, plan.get("blocks") or plan.get("sections") or [], api)
        settings = plan.get("settings") or {}
        if settings.get("update_fields_on_open", True):
            self._python_mark_fields_for_update(document, api)

    @staticmethod
    def _python_set_core_properties(document, metadata):
        props = document.core_properties
        mapping = {
            "title": "title", "subject": "subject", "author": "author",
            "keywords": "keywords", "comments": "comments", "category": "category",
        }
        for key, attr in mapping.items():
            if key in metadata:
                setattr(props, attr, str(metadata[key]))
        if not props.language:
            props.language = str(metadata.get("language") or "he-IL")

    def _python_configure_defaults(self, document, plan, api):
        Pt, RGBColor, qn = api["Pt"], api["RGBColor"], api["qn"]
        defaults = plan.get("defaults") or {}
        font_name = str(defaults.get("font") or "Arial")
        font_size = _clamp(defaults.get("font_size_pt"), 6, 72, 11)
        normal = document.styles["Normal"]
        normal.font.name = font_name
        normal.font.size = Pt(font_size)
        normal.font.color.rgb = RGBColor.from_string(_safe_color(defaults.get("color")))
        normal.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        normal.element.rPr.rFonts.set(qn("w:cs"), font_name)
        normal.paragraph_format.space_after = Pt(_clamp(defaults.get("space_after_pt"), 0, 72, 6))
        normal.paragraph_format.line_spacing = _clamp(defaults.get("line_spacing"), 0.8, 3, 1.15)
        for name, size, color in (
            ("Title", 26, "17365D"), ("Subtitle", 14, "4F6275"),
            ("Heading 1", 18, "17365D"), ("Heading 2", 15, "2F5597"),
            ("Heading 3", 12, "365F91"),
        ):
            try:
                style = document.styles[name]
                style.font.name = font_name
                style.font.size = Pt(size)
                style.font.color.rgb = RGBColor.from_string(color)
                style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
                style.element.rPr.rFonts.set(qn("w:cs"), font_name)
            except Exception:
                pass

    def _python_set_page_layout(self, document, page, api, sections=None):
        page = _merge_nested_spec(page, "page", "layout")
        Cm, WD_ORIENT = api["Cm"], api["WD_ORIENT"]
        OxmlElement, qn = api["OxmlElement"], api["qn"]
        size = str(page.get("size") or page.get("page_size") or "A4").upper()
        orientation = str(page.get("orientation") or "portrait").lower()
        margins = page.get("margins_cm") or {}
        for section in (sections if sections is not None else document.sections):
            if size == "A4":
                width, height = Cm(21.0), Cm(29.7)
            elif size in {"LETTER", "US LETTER"}:
                width, height = Cm(21.59), Cm(27.94)
            else:
                width = Cm(_clamp(page.get("width_cm"), 5, 100, 21.0))
                height = Cm(_clamp(page.get("height_cm"), 5, 100, 29.7))
            if orientation == "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
                width, height = height, width
            else:
                section.orientation = WD_ORIENT.PORTRAIT
            section.page_width, section.page_height = width, height
            section.top_margin = Cm(_clamp(margins.get("top"), 0, 10, 2.5))
            section.bottom_margin = Cm(_clamp(margins.get("bottom"), 0, 10, 2.5))
            section.left_margin = Cm(_clamp(margins.get("left"), 0, 10, 2.0))
            section.right_margin = Cm(_clamp(margins.get("right"), 0, 10, 2.0))
            section.gutter = Cm(_clamp(page.get("gutter_cm"), 0, 10, 0))
            section.header_distance = Cm(_clamp(page.get("header_distance_cm"), 0, 10, 1.25))
            section.footer_distance = Cm(_clamp(page.get("footer_distance_cm"), 0, 10, 1.25))
            section.different_first_page_header_footer = bool(page.get("different_first_page", False))
            bidi = section._sectPr.find(qn("w:bidi"))
            if page.get("rtl", True):
                if bidi is None:
                    bidi = OxmlElement("w:bidi")
                    section._sectPr.append(bidi)
                bidi.set(qn("w:val"), "1")
            elif bidi is not None:
                section._sectPr.remove(bidi)
        if "mirror_margins" in page:
            settings = document.settings.element
            mirror = settings.find(qn("w:mirrorMargins"))
            if page.get("mirror_margins") and mirror is None:
                settings.append(OxmlElement("w:mirrorMargins"))
            elif not page.get("mirror_margins") and mirror is not None:
                settings.remove(mirror)

    def _python_define_style(self, document, spec, api):
        spec = _normalized_format_spec(_merge_nested_spec(spec, "style"))
        WD_STYLE_TYPE = api["WD_STYLE_TYPE"]
        name = str(spec.get("name") or "").strip()
        if not name:
            raise ValueError("Style requires a name.")
        try:
            style = document.styles[name]
        except KeyError:
            style_type = WD_STYLE_TYPE.CHARACTER if str(spec.get("type") or "paragraph").lower() == "character" else WD_STYLE_TYPE.PARAGRAPH
            style = document.styles.add_style(name, style_type)
        self._python_apply_font(style.font, spec, api)
        if hasattr(style, "paragraph_format"):
            self._python_apply_paragraph_format(style.paragraph_format, spec, api)
        if spec.get("based_on"):
            try:
                style.base_style = document.styles[str(spec["based_on"])]
            except Exception:
                pass

    def _python_add_blocks(self, container, blocks, api):
        if not isinstance(blocks, list):
            raise ValueError("document.blocks must be an array.")
        for block in blocks:
            if not isinstance(block, dict):
                continue
            kind = str(block.get("type") or "paragraph").lower()
            if kind not in _PYTHON_BLOCK_TYPES:
                raise ValueError(f"Block type '{kind}' requires engine='com'.")
            if kind in {"paragraph", "heading", "title", "subtitle", "quote", "callout", "hyperlink", "bookmark", "field"}:
                self._python_add_text_block(container, block, kind, api)
            elif kind == "list":
                self._python_add_list(container, block, api)
            elif kind == "table":
                self._python_add_table(container, block, api)
            elif kind == "image":
                self._python_add_image(container, block, api)
            elif kind == "page_break":
                container.add_page_break()
            elif kind == "section_break":
                start = str(block.get("start") or "new_page").lower()
                enum = {
                    "continuous": api["WD_SECTION_START"].CONTINUOUS,
                    "even_page": api["WD_SECTION_START"].EVEN_PAGE,
                    "odd_page": api["WD_SECTION_START"].ODD_PAGE,
                }.get(start, api["WD_SECTION_START"].NEW_PAGE)
                section = container.add_section(enum)
                if block.get("page"):
                    self._python_set_page_layout(container, block.get("page") or {}, api, sections=[section])
            elif kind == "toc":
                self._python_ensure_toc_styles(container, api)
                paragraph = container.add_paragraph()
                self._python_set_paragraph(paragraph, block, api)
                self._python_add_field(paragraph, str(block.get("code") or 'TOC \\o "1-3" \\h \\z \\u'), "Table of Contents", api)
            elif kind in {"header", "footer"}:
                self._python_set_header_footer(container, kind, block, api)
            elif kind == "content_control":
                paragraph = container.add_paragraph()
                self._python_add_content_control(paragraph, block, api)
                self._python_set_paragraph(paragraph, block, api)

    def _python_add_text_block(self, container, block, kind, api):
        style = block.get("style")
        if not style:
            if kind == "heading":
                style = f"Heading {int(_clamp(block.get('level'), 1, 9, 1))}"
            elif kind == "title":
                style = "Title"
            elif kind == "subtitle":
                style = "Subtitle"
            elif kind == "quote":
                style = "Intense Quote"
        try:
            paragraph = container.add_paragraph(style=style) if style else container.add_paragraph()
        except Exception:
            paragraph = container.add_paragraph()
        runs = block.get("runs")
        if isinstance(runs, list):
            for run_spec in runs:
                if not isinstance(run_spec, dict):
                    run_spec = {"text": str(run_spec)}
                if run_spec.get("hyperlink"):
                    self._python_add_hyperlink(
                        paragraph,
                        str(run_spec.get("text") or ""),
                        str(run_spec["hyperlink"]),
                        run_spec,
                        api,
                    )
                elif run_spec.get("field"):
                    self._python_add_field(
                        paragraph,
                        self._word_field_code({"code": run_spec["field"]}),
                        str(run_spec.get("text") or ""),
                        api,
                    )
                else:
                    run = paragraph.add_run(str(run_spec.get("text") or ""))
                    self._python_apply_run(run, run_spec, api)
        elif kind == "hyperlink":
            link_text = str(block.get("text") or block.get("url") or block.get("anchor") or "")
            self._python_add_hyperlink(paragraph, link_text, str(block.get("url") or ""), block, api)
        elif kind == "field":
            self._python_add_field(paragraph, self._word_field_code(block), str(block.get("text") or ""), api)
        else:
            run = paragraph.add_run(str(block.get("text") or ""))
            self._python_apply_run(run, block, api)
        if kind == "bookmark":
            self._python_wrap_bookmark(paragraph, str(block.get("name") or "bookmark"), api)
        self._python_set_paragraph(paragraph, block, api)
        return paragraph

    def _python_ensure_toc_styles(self, document, api):
        """Create RTL TOC styles before Word materializes field results."""
        WD_STYLE_TYPE, OxmlElement, qn = api["WD_STYLE_TYPE"], api["OxmlElement"], api["qn"]
        section = document.sections[0]
        content_width_twips = max(
            720,
            int((section.page_width - section.left_margin - section.right_margin) / 635),
        )
        normal = document.styles["Normal"]
        for level in range(1, 10):
            name = f"TOC {level}"
            try:
                style = document.styles[name]
            except KeyError:
                style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = normal
            style.font.name = normal.font.name or "Arial"
            style.font.size = normal.font.size
            p_pr = style.element.get_or_add_pPr()
            bidi = p_pr.find(qn("w:bidi"))
            if bidi is None:
                bidi = OxmlElement("w:bidi")
                p_pr.append(bidi)
            bidi.set(qn("w:val"), "1")
            jc = p_pr.find(qn("w:jc"))
            if jc is None:
                jc = OxmlElement("w:jc")
                p_pr.append(jc)
            jc.set(qn("w:val"), "left")  # physical right after bidi mirroring
            tabs = p_pr.find(qn("w:tabs"))
            if tabs is None:
                tabs = OxmlElement("w:tabs")
                p_pr.append(tabs)
            for existing in list(tabs):
                tabs.remove(existing)
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "right")
            tab.set(qn("w:leader"), "dot")
            tab.set(qn("w:pos"), str(content_width_twips))
            tabs.append(tab)
            if level > 1:
                style.paragraph_format.left_indent = api["Cm"](0.45 * (level - 1))

    def _python_set_paragraph(self, paragraph, spec, api):
        spec = _normalized_format_spec(spec)
        WD_ALIGN_PARAGRAPH, OxmlElement, qn = api["WD_ALIGN_PARAGRAPH"], api["OxmlElement"], api["qn"]
        if spec.get("style"):
            try:
                paragraph.style = str(spec["style"])
            except (KeyError, ValueError):
                pass
        rtl = bool(spec.get("rtl", True))
        alignment = str(spec.get("alignment") or ("right" if rtl else "left")).lower()
        # In WordprocessingML, Word mirrors the physical left/right meaning of
        # w:jc for bidi paragraphs.  Swap only those two values so callers keep
        # using visual alignment names ("right" means the right page margin).
        xml_alignment = {"left": "right", "right": "left"}.get(alignment, alignment) if rtl else alignment
        paragraph.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }.get(xml_alignment, WD_ALIGN_PARAGRAPH.LEFT if rtl else WD_ALIGN_PARAGRAPH.RIGHT)
        self._python_apply_paragraph_format(paragraph.paragraph_format, spec, api)
        p_pr = paragraph._p.get_or_add_pPr()
        bidi = p_pr.find(qn("w:bidi"))
        if rtl:
            if bidi is None:
                bidi = OxmlElement("w:bidi")
                p_pr.append(bidi)
            bidi.set(qn("w:val"), "1")
        elif bidi is not None:
            p_pr.remove(bidi)
        if spec.get("keep_with_next"):
            keep = OxmlElement("w:keepNext")
            p_pr.append(keep)
        if spec.get("page_break_before"):
            p_pr.append(OxmlElement("w:pageBreakBefore"))

    def _python_apply_paragraph_format(self, fmt, spec, api):
        spec = _normalized_format_spec(spec)
        Pt, Cm = api["Pt"], api["Cm"]
        for key, attr in (
            ("space_before_pt", "space_before"), ("space_after_pt", "space_after"),
        ):
            if key in spec:
                setattr(fmt, attr, Pt(_clamp(spec[key], 0, 300, 0)))
        if "line_spacing" in spec:
            fmt.line_spacing = _clamp(spec["line_spacing"], 0.5, 5, 1.15)
        for key, attr in (
            ("left_indent_cm", "left_indent"), ("right_indent_cm", "right_indent"),
            ("first_line_indent_cm", "first_line_indent"),
        ):
            if key in spec:
                setattr(fmt, attr, Cm(_clamp(spec[key], -10, 30, 0)))
        if "keep_together" in spec:
            fmt.keep_together = bool(spec["keep_together"])
        if "widow_control" in spec:
            fmt.widow_control = bool(spec["widow_control"])

    def _python_apply_font(self, font, spec, api):
        spec = _normalized_format_spec(spec)
        Pt, RGBColor = api["Pt"], api["RGBColor"]
        if spec.get("font"):
            font.name = str(spec["font"])
        if spec.get("font_size_pt") is not None:
            font.size = Pt(_clamp(spec["font_size_pt"], 1, 400, 11))
        for key, attr in (
            ("bold", "bold"), ("italic", "italic"), ("underline", "underline"),
            ("strike", "strike"), ("superscript", "superscript"), ("subscript", "subscript"),
        ):
            if key in spec:
                setattr(font, attr, bool(spec[key]))
        if spec.get("color"):
            font.color.rgb = RGBColor.from_string(_safe_color(spec["color"]))

    def _python_apply_run(self, run, spec, api):
        OxmlElement, qn = api["OxmlElement"], api["qn"]
        self._python_apply_font(run.font, spec, api)
        font_name = str(spec.get("font") or run.font.name or "Arial")
        r_pr = run._r.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:cs"), font_name)
        r_fonts.set(qn("w:eastAsia"), font_name)
        lang = r_pr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            r_pr.append(lang)
        lang.set(qn("w:val"), str(spec.get("language") or "he-IL"))
        lang.set(qn("w:bidi"), str(spec.get("language") or "he-IL"))
        if spec.get("rtl", True):
            rtl = OxmlElement("w:rtl")
            rtl.set(qn("w:val"), "1")
            r_pr.append(rtl)

    def _python_add_hyperlink(self, paragraph, text, url, spec, api):
        url, anchor = self._hyperlink_target(url, spec.get("anchor"))
        OxmlElement, qn = api["OxmlElement"], api["qn"]
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:history"), "1")
        if anchor:
            hyperlink.set(qn("w:anchor"), anchor)
        else:
            relationship = paragraph.part.relate_to(
                url,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True,
            )
            hyperlink.set(qn("r:id"), relationship)
        run = paragraph.add_run(text)
        self._python_apply_run(run, spec, api)
        r_pr = run._r.get_or_add_rPr()
        if not spec.get("color"):
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "0563C1")
            r_pr.append(color)
        if "underline" not in spec:
            underline = OxmlElement("w:u")
            underline.set(qn("w:val"), "single")
            r_pr.append(underline)
        hyperlink.append(run._r)
        paragraph._p.append(hyperlink)

    def _python_add_list(self, container, block, api):
        ordered = bool(block.get("ordered", False))
        items = block.get("items") or []
        if not isinstance(items, list) or not items:
            raise ValueError("list.items must be a non-empty array.")
        for raw_item in items:
            item = raw_item if isinstance(raw_item, dict) else {"text": str(raw_item)}
            level = int(_clamp(item.get("level"), 0, 8, 0))
            base_style = "List Number" if item.get("ordered", ordered) else "List Bullet"
            style = item.get("style") or (base_style if level == 0 else f"{base_style} {min(level + 1, 3)}")
            paragraph = self._python_add_text_block(container, {**block, **item, "style": style}, "paragraph", api)
            if level > 2:
                paragraph.paragraph_format.left_indent = api["Cm"](0.63 * (level - 2))

    @staticmethod
    def _python_add_field(paragraph, code, display_text, api):
        OxmlElement, qn = api["OxmlElement"], api["qn"]
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = code
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        text_node = OxmlElement("w:t")
        text_node.text = display_text
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        for element in (begin, instruction, separate, text_node, end):
            run = OxmlElement("w:r")
            run.append(element)
            paragraph._p.append(run)

    @staticmethod
    def _python_wrap_bookmark(paragraph, name, api):
        OxmlElement, qn = api["OxmlElement"], api["qn"]
        clean = _bookmark_name(name)
        bookmark_id = str(abs(hash((clean, paragraph.text))) % 2000000000)
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), bookmark_id)
        start.set(qn("w:name"), clean)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), bookmark_id)
        paragraph._p.insert(0, start)
        paragraph._p.append(end)

    def _python_add_table(self, container, block, api):
        Cm, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT = api["Cm"], api["WD_CELL_VERTICAL_ALIGNMENT"], api["WD_TABLE_ALIGNMENT"]
        OxmlElement, qn = api["OxmlElement"], api["qn"]
        rows = block.get("rows") or []
        if not isinstance(rows, list) or not rows:
            raise ValueError("table.rows must be a non-empty array.")
        normalized = [row if isinstance(row, list) else list(row.values()) if isinstance(row, dict) else [row] for row in rows]
        columns = max(len(row) for row in normalized)
        table = container.add_table(rows=len(normalized), cols=columns)
        try:
            table.style = str(block.get("style") or "Table Grid")
        except Exception:
            pass
        table.alignment = {
            "left": WD_TABLE_ALIGNMENT.LEFT, "center": WD_TABLE_ALIGNMENT.CENTER,
            "right": WD_TABLE_ALIGNMENT.RIGHT,
        }.get(str(block.get("alignment") or "right").lower(), WD_TABLE_ALIGNMENT.RIGHT)
        table.autofit = bool(block.get("autofit", False))
        widths = block.get("column_widths_cm") or []
        header_rows = int(_clamp(block.get("header_rows"), 0, len(normalized), 1))
        for row_index, values in enumerate(normalized):
            for col_index in range(columns):
                cell = table.cell(row_index, col_index)
                value = values[col_index] if col_index < len(values) else ""
                spec = value if isinstance(value, dict) else {"text": value}
                cell.text = ""
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(str(spec.get("text") or ""))
                merged = {**block.get("cell_defaults", {}), **spec}
                if row_index < header_rows:
                    merged.setdefault("bold", True)
                self._python_apply_run(run, merged, api)
                self._python_set_paragraph(paragraph, merged, api)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if col_index < len(widths):
                    cell.width = Cm(_clamp(widths[col_index], 0.5, 50, 3))
                tc_pr = cell._tc.get_or_add_tcPr()
                if merged.get("fill"):
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), _safe_color(merged["fill"], "FFFFFF"))
                    tc_pr.append(shading)
            if row_index < header_rows:
                tr_pr = table.rows[row_index]._tr.get_or_add_trPr()
                tr_pr.append(OxmlElement("w:tblHeader"))
        tbl_pr = table._tbl.tblPr
        bidi_visual = OxmlElement("w:bidiVisual")
        bidi_visual.set(qn("w:val"), "1" if block.get("rtl", True) else "0")
        tbl_pr.append(bidi_visual)
        for merge in block.get("merges") or []:
            if not isinstance(merge, dict):
                continue
            start_row = int(merge.get("from_row", merge.get("row", 0)))
            start_col = int(merge.get("from_column", merge.get("column", 0)))
            end_row = int(merge.get("to_row", start_row))
            end_col = int(merge.get("to_column", start_col))
            if not (0 <= start_row <= end_row < len(normalized) and 0 <= start_col <= end_col < columns):
                raise ValueError("table merge coordinates are out of range.")
            table.cell(start_row, start_col).merge(table.cell(end_row, end_col))
        return table

    def _python_add_image(self, container, block, api):
        path = self._document_resolve_path(block.get("path"), mode="read")
        Cm, WD_ALIGN_PARAGRAPH = api["Cm"], api["WD_ALIGN_PARAGRAPH"]
        paragraph = container.add_paragraph()
        paragraph.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }.get(str(block.get("alignment") or "center").lower(), WD_ALIGN_PARAGRAPH.CENTER)
        run = paragraph.add_run()
        kwargs = {}
        if block.get("width_cm") is not None:
            kwargs["width"] = Cm(_clamp(block["width_cm"], 0.1, 100, 10))
        if block.get("height_cm") is not None:
            kwargs["height"] = Cm(_clamp(block["height_cm"], 0.1, 100, 10))
        shape = run.add_picture(path, **kwargs)
        if block.get("alt_text"):
            doc_pr = shape._inline.docPr
            doc_pr.set("descr", str(block["alt_text"])[:1000])
            doc_pr.set("title", str(block.get("title") or block["alt_text"])[:255])
        return paragraph

    def _python_set_header_footer(self, document, kind, spec, api):
        spec = _normalized_format_spec(_merge_nested_spec(spec, "content"))
        for section in document.sections:
            if "different_first_page" in spec:
                section.different_first_page_header_footer = bool(spec["different_first_page"])
            story = section.header if kind == "header" else section.footer
            paragraph = story.paragraphs[0]
            paragraph.clear()
            if spec.get("runs"):
                for item in spec["runs"]:
                    item = item if isinstance(item, dict) else {"text": str(item)}
                    if item.get("field"):
                        self._python_add_field(
                            paragraph,
                            self._word_field_code({"code": item["field"]}),
                            str(item.get("text") or ""),
                            api,
                        )
                    else:
                        run = paragraph.add_run(str(item.get("text") or ""))
                        self._python_apply_run(run, item, api)
            else:
                run = paragraph.add_run(str(spec.get("text") or ""))
                self._python_apply_run(run, spec, api)
                if spec.get("page_number"):
                    paragraph.add_run(" ")
                    self._python_add_field(paragraph, "PAGE", "1", api)
            self._python_set_paragraph(paragraph, spec, api)

    def _python_add_content_control(self, paragraph, spec, api):
        OxmlElement, qn = api["OxmlElement"], api["qn"]
        control_type = _friendly_enum(
            spec.get("control_type"), _WORD_CONTENT_CONTROL_TYPES, "control_type", 0
        )
        sdt = OxmlElement("w:sdt")
        sdt_pr = OxmlElement("w:sdtPr")
        alias = OxmlElement("w:alias")
        alias.set(qn("w:val"), str(spec.get("title") or "Content"))
        tag = OxmlElement("w:tag")
        tag.set(qn("w:val"), str(spec.get("tag") or "smarti"))
        sdt_pr.extend([alias, tag])
        if control_type == 1:
            sdt_pr.append(OxmlElement("w:text"))
        elif control_type == 2:
            sdt_pr.append(OxmlElement("w:picture"))
        elif control_type in {3, 4}:
            chooser = OxmlElement("w:comboBox" if control_type == 3 else "w:dropDownList")
            for raw_item in spec.get("items") or []:
                item = raw_item if isinstance(raw_item, dict) else {"display": raw_item, "value": raw_item}
                option = OxmlElement("w:listItem")
                option.set(qn("w:displayText"), str(item.get("display") or item.get("text") or item.get("value") or ""))
                option.set(qn("w:value"), str(item.get("value") or item.get("display") or item.get("text") or ""))
                chooser.append(option)
            sdt_pr.append(chooser)
        elif control_type == 6:
            date = OxmlElement("w:date")
            date_format = OxmlElement("w:dateFormat")
            date_format.set(qn("w:val"), str(spec.get("date_format") or "dd/MM/yyyy"))
            date.append(date_format)
            sdt_pr.append(date)
        elif control_type == 8:
            checkbox = OxmlElement("w14:checkbox")
            checked = OxmlElement("w14:checked")
            checked.set(qn("w14:val"), "1" if spec.get("checked") else "0")
            checkbox.append(checked)
            sdt_pr.append(checkbox)
        content = OxmlElement("w:sdtContent")
        default_text = "☒" if spec.get("checked") else "☐" if control_type == 8 else ""
        run = paragraph.add_run(str(spec.get("text") or spec.get("placeholder") or default_text))
        self._python_apply_run(run, spec, api)
        content.append(run._r)
        sdt.extend([sdt_pr, content])
        paragraph._p.append(sdt)

    @staticmethod
    def _python_mark_fields_for_update(document, api):
        OxmlElement, qn = api["OxmlElement"], api["qn"]
        settings = document.settings.element
        existing = settings.find(qn("w:updateFields"))
        if existing is None:
            existing = OxmlElement("w:updateFields")
            settings.append(existing)
        existing.set(qn("w:val"), "true")

    def _python_edit_document(self, source, output, operations):
        api = self._python_imports()
        document = api["docx"].Document(source)
        counts = {}
        for operation in operations:
            if not isinstance(operation, dict):
                continue
            op = str(operation.get("op") or operation.get("type") or "").lower()
            if op in _COM_ONLY_OPERATIONS:
                raise ValueError(f"Operation '{op}' requires engine='com'.")
            if op == "replace_text":
                old = str(operation.get("find") or operation.get("old") or "")
                new = str(operation.get("replace") if "replace" in operation else operation.get("new") or "")
                if not old:
                    raise ValueError("replace_text requires find.")
                counts[op] = counts.get(op, 0) + self._python_replace_text(
                    document,
                    old,
                    new,
                    bool(operation.get("case_sensitive")),
                    bool(operation.get("whole_word")),
                    bool(operation.get("replace_all", True)),
                )
            elif op in {"append_blocks", "append"}:
                blocks = operation.get("blocks") or []
                self._python_add_blocks(document, blocks, api)
                counts[op] = counts.get(op, 0) + len(blocks)
            elif op in {"insert_blocks", "insert"}:
                blocks = operation.get("blocks") or []
                paragraph = self._python_select_paragraph(document, operation.get("selector") or {})
                self._python_insert_blocks(
                    document,
                    paragraph,
                    blocks,
                    str(operation.get("position") or "before").lower(),
                    api,
                )
                counts[op] = counts.get(op, 0) + len(blocks)
            elif op == "delete_paragraph":
                paragraph = self._python_select_paragraph(document, operation.get("selector") or operation)
                paragraph._element.getparent().remove(paragraph._element)
                counts[op] = counts.get(op, 0) + 1
            elif op == "format_paragraph":
                paragraph = self._python_select_paragraph(document, operation.get("selector") or operation)
                self._python_set_paragraph(paragraph, operation.get("format") or operation, api)
                for run in paragraph.runs:
                    self._python_apply_run(run, operation.get("format") or operation, api)
                counts[op] = counts.get(op, 0) + 1
            elif op == "set_page_layout":
                self._python_set_page_layout(document, operation, api)
                counts[op] = counts.get(op, 0) + 1
            elif op == "define_style":
                self._python_define_style(document, operation, api)
                counts[op] = counts.get(op, 0) + 1
            elif op in {"set_header", "set_footer"}:
                self._python_set_header_footer(document, op.split("_")[1], operation, api)
                counts[op] = counts.get(op, 0) + 1
            elif op == "update_fields":
                self._python_mark_fields_for_update(document, api)
                counts[op] = counts.get(op, 0) + 1
            elif op == "set_properties":
                self._python_set_core_properties(document, operation.get("properties") or operation)
                counts[op] = counts.get(op, 0) + 1
            else:
                raise ValueError(f"Unsupported python edit operation: {op}")
        temp = self._document_temp_path(output)
        try:
            document.save(temp)
            os.replace(temp, output)
        finally:
            if os.path.exists(temp):
                os.remove(temp)
        return {"status": "edited", "operation_counts": counts}

    def _python_insert_blocks(self, document, paragraph, blocks, position, api):
        """Build portable blocks with python-docx, then move them beside an anchor."""
        marker = document.add_paragraph()._p
        self._python_add_blocks(document, blocks, api)
        qn = api["qn"]
        added = []
        node = marker.getnext()
        while node is not None and node.tag != qn("w:sectPr"):
            following = node.getnext()
            added.append(node)
            node = following
        marker.getparent().remove(marker)

        anchor = paragraph._p
        if position == "before":
            for node in added:
                anchor.addprevious(node)
        elif position == "after":
            for node in added:
                anchor.addnext(node)
                anchor = node
        else:
            raise ValueError("insert_blocks position must be 'before' or 'after'.")

    def _python_iter_paragraphs(self, document):
        for paragraph in document.paragraphs:
            yield paragraph
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph
        for section in document.sections:
            for story in (section.header, section.footer):
                for paragraph in story.paragraphs:
                    yield paragraph

    def _python_replace_text(self, document, old, new, case_sensitive=False, whole_word=False, replace_all=True):
        count = 0
        escaped = re.escape(old)
        if whole_word:
            escaped = rf"(?<!\w){escaped}(?!\w)"
        pattern = re.compile(escaped, 0 if case_sensitive else re.IGNORECASE)
        for paragraph in self._python_iter_paragraphs(document):
            runs = list(paragraph.runs)
            combined = "".join(run.text for run in runs)
            matches = list(pattern.finditer(combined))
            if matches and not replace_all:
                matches = matches[:1]
            if not matches:
                continue
            starts = []
            cursor = 0
            for run in runs:
                starts.append(cursor)
                cursor += len(run.text)
            for match in reversed(matches):
                start, end = match.span()
                start_index = next(index for index, offset in enumerate(starts) if offset + len(runs[index].text) > start)
                end_index = next(index for index, offset in enumerate(starts) if offset + len(runs[index].text) >= end)
                start_offset = start - starts[start_index]
                end_offset = end - starts[end_index]
                if start_index == end_index:
                    text = runs[start_index].text
                    runs[start_index].text = text[:start_offset] + new + text[end_offset:]
                else:
                    runs[start_index].text = runs[start_index].text[:start_offset] + new
                    for index in range(start_index + 1, end_index):
                        runs[index].text = ""
                    runs[end_index].text = runs[end_index].text[end_offset:]
            count += len(matches)
            if count and not replace_all:
                break
        return count

    @staticmethod
    def _python_select_paragraph(document, selector):
        selector = selector if isinstance(selector, dict) else {}
        if selector.get("paragraph_index") is not None:
            index = int(selector["paragraph_index"])
            if index < 0 or index >= len(document.paragraphs):
                raise IndexError("paragraph_index is out of range.")
            return document.paragraphs[index]
        text = str(selector.get("find") or selector.get("contains") or "")
        occurrence = max(1, int(selector.get("occurrence") or 1))
        seen = 0
        for paragraph in document.paragraphs:
            if text in paragraph.text:
                seen += 1
                if seen == occurrence:
                    return paragraph
        raise ValueError("Paragraph selector did not match.")

    # -------------------------- Word COM engine ---------------------------

    def _com_create_document(self, output, template, plan, args):
        temp = self._document_temp_path(output)
        try:
            with _WordComSession(args.get("timeout_seconds", 180), args.get("visible", False)) as session:
                document = session.new_document(template)
                self._com_apply_plan(document, plan, args)
                self._com_save_document(document, temp, output)
                session.close_document(False)
            os.replace(temp, output)
        finally:
            if os.path.exists(temp):
                os.remove(temp)
        return {"status": "created", "blocks": len(plan.get("blocks") or plan.get("sections") or [])}

    def _com_edit_document(self, source, output, operations, args):
        temp = self._document_temp_path(output)
        results = []
        try:
            with _WordComSession(args.get("timeout_seconds", 180), args.get("visible", False)) as session:
                document = session.open_document(source, password=str(args.get("password") or ""))
                for operation in operations:
                    if isinstance(operation, dict):
                        results.append(self._com_apply_operation(document, operation, args))
                try:
                    self._com_prepare_toc_styles(document)
                    document.Fields.Update()
                except Exception:
                    pass
                self._com_save_document(document, temp, output)
                session.close_document(False)
            os.replace(temp, output)
        finally:
            if os.path.exists(temp):
                os.remove(temp)
        return {"status": "edited", "operations": results}

    def _com_apply_plan(self, document, plan, args):
        self._com_set_properties(document, plan.get("metadata") or {})
        self._com_set_page_layout(document, plan.get("page") or {})
        defaults = plan.get("defaults") or {}
        self._com_configure_styles(document, defaults)
        for style in plan.get("styles") or []:
            if isinstance(style, dict):
                self._com_define_style(document, style)
        if isinstance(plan.get("header"), dict):
            self._com_set_header_footer(document, "header", plan["header"])
        if isinstance(plan.get("footer"), dict):
            self._com_set_header_footer(document, "footer", plan["footer"])
        self._com_add_blocks(document, plan.get("blocks") or plan.get("sections") or [], args)
        if (plan.get("settings") or {}).get("track_changes"):
            document.TrackRevisions = True
        try:
            self._com_prepare_toc_styles(document)
            document.Fields.Update()
        except Exception:
            pass

    @staticmethod
    def _com_set_properties(document, metadata):
        mapping = {
            "title": "Title", "subject": "Subject", "author": "Author",
            "keywords": "Keywords", "comments": "Comments", "category": "Category",
        }
        for key, prop_name in mapping.items():
            if key not in metadata:
                continue
            try:
                document.BuiltInDocumentProperties(prop_name).Value = str(metadata[key])
            except Exception:
                pass

    def _com_set_page_layout(self, document, page, sections=None):
        page = _merge_nested_spec(page, "page", "layout")
        size = str(page.get("size") or page.get("page_size") or "A4").upper()
        orientation = str(page.get("orientation") or "portrait").lower()
        margins = page.get("margins_cm") or {}
        selected_sections = sections or [document.Sections(index) for index in range(1, int(document.Sections.Count) + 1)]
        for section in selected_sections:
            setup = section.PageSetup
            setup.PaperSize = 7 if size == "A4" else 2 if size in {"LETTER", "US LETTER"} else setup.PaperSize
            setup.Orientation = 1 if orientation == "landscape" else 0
            setup.TopMargin = self._cm_to_points(margins.get("top", 2.5))
            setup.BottomMargin = self._cm_to_points(margins.get("bottom", 2.5))
            setup.LeftMargin = self._cm_to_points(margins.get("left", 2.0))
            setup.RightMargin = self._cm_to_points(margins.get("right", 2.0))
            setup.Gutter = self._cm_to_points(page.get("gutter_cm", 0))
            setup.HeaderDistance = self._cm_to_points(page.get("header_distance_cm", 1.25))
            setup.FooterDistance = self._cm_to_points(page.get("footer_distance_cm", 1.25))
            setup.DifferentFirstPageHeaderFooter = bool(page.get("different_first_page", False))
            setup.OddAndEvenPagesHeaderFooter = bool(page.get("different_odd_even", False))
            if "mirror_margins" in page:
                setup.MirrorMargins = bool(page["mirror_margins"])
            try:
                setup.SectionDirection = 0 if page.get("rtl", True) else 1
            except Exception:
                pass

    @staticmethod
    def _cm_to_points(value):
        return _clamp(value, -100, 1000, 0) * 72.0 / 2.54

    def _com_configure_styles(self, document, defaults):
        font = str(defaults.get("font") or "Arial")
        size = _clamp(defaults.get("font_size_pt"), 6, 72, 11)
        for style_ref, font_size, bold, color in (
            (-1, size, False, "000000"), (-63, 26, True, "17365D"),
            (-75, 14, False, "4F6275"), (-2, 18, True, "17365D"),
            (-3, 15, True, "2F5597"), (-4, 12, True, "365F91"),
        ):
            try:
                style = document.Styles(style_ref)
                style.Font.Name = font
                style.Font.NameBi = font
                style.Font.Size = font_size
                style.Font.SizeBi = font_size
                style.Font.Bold = bold
                style.Font.Color = self._com_rgb(color)
                style.ParagraphFormat.ReadingOrder = 0
            except Exception:
                pass
        self._com_prepare_toc_styles(document)

    @staticmethod
    def _com_prepare_toc_styles(document):
        try:
            rtl = int(document.Sections(1).PageSetup.SectionDirection) == 0
        except Exception:
            rtl = False
        if not rtl:
            return
        for style_ref in range(-20, -29, -1):  # wdStyleTOC1 through wdStyleTOC9
            try:
                style = document.Styles(style_ref)
                style.ParagraphFormat.ReadingOrder = 0
            except Exception:
                pass

    def _com_define_style(self, document, spec):
        spec = _normalized_format_spec(_merge_nested_spec(spec, "style"))
        name = str(spec.get("name") or "").strip()
        if not name:
            raise ValueError("Style requires a name.")
        try:
            style = document.Styles(self._com_style_ref(name))
        except Exception:
            style = document.Styles.Add(Name=name, Type=1 if str(spec.get("type") or "paragraph") == "paragraph" else 2)
        self._com_apply_font(style.Font, spec)
        # Setting Alignment on a Style.ParagraphFormat can make current Word
        # builds query/restart the printer-backed layout service. Paragraphs
        # receive direction/alignment safely when the style is applied.
        self._com_apply_paragraph_format(style.ParagraphFormat, spec, apply_direction=False)
        if spec.get("based_on"):
            try:
                style.BaseStyle = self._com_style_ref(spec["based_on"])
            except Exception:
                pass

    @staticmethod
    def _com_style_ref(value):
        if isinstance(value, int):
            return value
        text = str(value or "").strip()
        folded = re.sub(r"[\s_-]+", "", text).casefold()
        return _WORD_BUILTIN_STYLES.get(folded, text)

    @staticmethod
    def _com_rgb(value):
        text = _safe_color(value)
        red, green, blue = int(text[0:2], 16), int(text[2:4], 16), int(text[4:6], 16)
        return red | (green << 8) | (blue << 16)

    def _com_apply_font(self, font, spec):
        spec = _normalized_format_spec(spec)
        if spec.get("font"):
            font.Name = str(spec["font"])
            font.NameBi = str(spec["font"])
        if spec.get("font_size_pt") is not None:
            size = _clamp(spec["font_size_pt"], 1, 400, 11)
            font.Size = size
            font.SizeBi = size
        for key, attr in (
            ("bold", "Bold"), ("italic", "Italic"), ("underline", "Underline"),
            ("strike", "StrikeThrough"), ("superscript", "Superscript"), ("subscript", "Subscript"),
        ):
            if key in spec:
                setattr(font, attr, bool(spec[key]))
        if spec.get("color"):
            font.Color = self._com_rgb(spec["color"])

    def _com_apply_paragraph_format(self, fmt, spec, apply_direction=True, range_obj=None):
        spec = _normalized_format_spec(spec)
        if apply_direction:
            alignment = str(spec.get("alignment") or ("right" if spec.get("rtl", True) else "left")).lower()
            fmt.ReadingOrder = 0 if spec.get("rtl", True) else 1
            if range_obj is None:
                fmt.Alignment = {"left": 0, "center": 1, "right": 2, "justify": 3}.get(alignment, 2)
            else:
                import pythoncom
                command = {
                    "left": "LeftPara", "center": "CenterPara",
                    "right": "RightPara", "justify": "JustifyPara",
                }.get(alignment, "RightPara")
                range_obj.Select()
                ole = range_obj.Application.WordBasic._oleobj_
                ole.Invoke(
                    ole.GetIDsOfNames(command), 0,
                    pythoncom.DISPATCH_METHOD, 1,
                )
        for key, attr in (
            ("space_before_pt", "SpaceBefore"), ("space_after_pt", "SpaceAfter"),
        ):
            if key in spec:
                setattr(fmt, attr, _clamp(spec[key], 0, 300, 0))
        if "line_spacing" in spec:
            fmt.LineSpacingRule = 5
            fmt.LineSpacing = 12 * _clamp(spec["line_spacing"], 0.5, 5, 1.15)
        for key, attr in (
            ("left_indent_cm", "LeftIndent"), ("right_indent_cm", "RightIndent"),
            ("first_line_indent_cm", "FirstLineIndent"),
        ):
            if key in spec:
                setattr(fmt, attr, self._cm_to_points(spec[key]))
        if "keep_with_next" in spec:
            fmt.KeepWithNext = bool(spec["keep_with_next"])
        if "keep_together" in spec:
            fmt.KeepTogether = bool(spec["keep_together"])
        if "widow_control" in spec:
            fmt.WidowControl = bool(spec["widow_control"])
        if "page_break_before" in spec:
            fmt.PageBreakBefore = bool(spec["page_break_before"])

    def _com_add_blocks(self, document, blocks, args, anchor=None):
        for block in blocks if isinstance(blocks, list) else []:
            if isinstance(block, dict):
                self._com_add_block(document, block, args, anchor)

    def _com_end_range(self, document):
        end = max(0, int(document.Content.End) - 1)
        return document.Range(end, end)

    def _com_add_block(self, document, block, args, anchor=None):
        kind = str(block.get("type") or "paragraph").lower()
        target = anchor.Duplicate if anchor is not None else self._com_end_range(document)
        try:
            target.Collapse(0)
        except Exception:
            pass
        if kind in {"paragraph", "heading", "title", "subtitle", "quote", "callout", "hyperlink", "bookmark", "field"}:
            start = int(target.Start)
            runs = block.get("runs")
            if isinstance(runs, list):
                for item in runs:
                    item = item if isinstance(item, dict) else {"text": str(item)}
                    run_start = int(target.End)
                    text = str(item.get("text") or "")
                    target.InsertAfter(text)
                    run_range = document.Range(run_start, run_start + len(text))
                    self._com_apply_font(run_range.Font, item)
                    run_range.LanguageID = 1037
                    if item.get("hyperlink"):
                        url, bookmark = self._hyperlink_target(item.get("hyperlink"), item.get("anchor"))
                        document.Hyperlinks.Add(
                            Anchor=run_range,
                            Address=url,
                            SubAddress=bookmark,
                            TextToDisplay=text,
                        )
                    elif item.get("field"):
                        document.Fields.Add(
                            Range=run_range,
                            Type=-1,
                            Text=self._word_field_code({"code": item.get("field")}),
                            PreserveFormatting=True,
                        )
            elif kind == "hyperlink":
                url, bookmark = self._hyperlink_target(block.get("url"), block.get("anchor"))
                document.Hyperlinks.Add(
                    Anchor=target,
                    Address=url,
                    SubAddress=bookmark,
                    TextToDisplay=str(block.get("text") or url or bookmark),
                )
            elif kind == "field":
                document.Fields.Add(
                    Range=target,
                    Type=-1,
                    Text=self._word_field_code(block),
                    PreserveFormatting=True,
                )
            else:
                text = str(block.get("text") or "")
                target.InsertAfter(text)
            end = int(target.End)
            paragraph_range = document.Range(start, end)
            paragraph = paragraph_range.Paragraphs(1)
            if kind == "heading":
                paragraph.Range.Style = -int(_clamp(block.get('level'), 1, 9, 1)) - 1
            elif kind == "title":
                paragraph.Range.Style = -63
            elif kind == "subtitle":
                paragraph.Range.Style = -75
            elif kind == "quote":
                paragraph.Range.Style = -181
            elif block.get("style"):
                paragraph.Range.Style = self._com_style_ref(block["style"])
            self._com_apply_font(paragraph.Range.Font, block)
            self._com_apply_paragraph_format(
                paragraph.Range.ParagraphFormat, block, range_obj=paragraph.Range
            )
            paragraph.Range.LanguageID = 1037
            if kind == "bookmark":
                document.Bookmarks.Add(_bookmark_name(block.get("name")), paragraph_range)
            paragraph.Range.InsertParagraphAfter()
            return
        if kind == "list":
            items = block.get("items") or []
            if not isinstance(items, list) or not items:
                raise ValueError("list.items must be a non-empty array.")
            start = int(target.Start)
            texts = [str(item.get("text") or "") if isinstance(item, dict) else str(item) for item in items]
            target.InsertAfter("\r".join(texts))
            list_range = document.Range(start, start + len("\r".join(texts)))
            if block.get("ordered", False):
                list_range.ListFormat.ApplyNumberDefault()
            else:
                list_range.ListFormat.ApplyBulletDefault()
            for index, raw_item in enumerate(items, 1):
                item = raw_item if isinstance(raw_item, dict) else {"text": str(raw_item)}
                paragraph = list_range.Paragraphs(index)
                merged = {**block, **item}
                self._com_apply_font(paragraph.Range.Font, merged)
                self._com_apply_paragraph_format(
                    paragraph.Format, merged, range_obj=paragraph.Range
                )
                paragraph.Range.LanguageID = 1037
                for _level in range(int(_clamp(item.get("level"), 0, 8, 0))):
                    paragraph.Range.ListFormat.ListIndent()
            list_range.InsertParagraphAfter()
            return
        if kind == "table":
            rows = block.get("rows") or []
            normalized = [row if isinstance(row, list) else list(row.values()) if isinstance(row, dict) else [row] for row in rows]
            if not normalized:
                raise ValueError("table.rows must be non-empty.")
            cols = max(len(row) for row in normalized)
            table = document.Tables.Add(target, len(normalized), cols)
            try:
                table.Style = str(block.get("style") or "Table Grid")
            except Exception:
                pass
            table.TableDirection = 0 if block.get("rtl", True) else 1
            table.Rows.Alignment = {"left": 0, "center": 1, "right": 2}.get(str(block.get("alignment") or "right").lower(), 2)
            widths = block.get("column_widths_cm") or []
            header_rows = int(_clamp(block.get("header_rows"), 0, len(normalized), 1))
            for row_index, values in enumerate(normalized, 1):
                for col_index in range(1, cols + 1):
                    value = values[col_index - 1] if col_index <= len(values) else ""
                    spec = value if isinstance(value, dict) else {"text": value}
                    cell = table.Cell(row_index, col_index)
                    cell.Range.Text = str(spec.get("text") or "")
                    merged = {**block.get("cell_defaults", {}), **spec}
                    if row_index <= header_rows:
                        merged.setdefault("bold", True)
                    self._com_apply_font(cell.Range.Font, merged)
                    self._com_apply_paragraph_format(
                        cell.Range.ParagraphFormat, merged, range_obj=cell.Range
                    )
                    cell.VerticalAlignment = 1
                    if col_index <= len(widths):
                        cell.Width = self._cm_to_points(widths[col_index - 1])
                    if merged.get("fill"):
                        cell.Shading.BackgroundPatternColor = self._com_rgb(merged["fill"])
                if row_index <= header_rows:
                    table.Rows(row_index).HeadingFormat = True
            for merge in block.get("merges") or []:
                if not isinstance(merge, dict):
                    continue
                start_row = int(merge.get("from_row", merge.get("row", 0))) + 1
                start_col = int(merge.get("from_column", merge.get("column", 0))) + 1
                end_row = int(merge.get("to_row", start_row - 1)) + 1
                end_col = int(merge.get("to_column", start_col - 1)) + 1
                if not (1 <= start_row <= end_row <= len(normalized) and 1 <= start_col <= end_col <= cols):
                    raise ValueError("table merge coordinates are out of range.")
                table.Cell(start_row, start_col).Merge(table.Cell(end_row, end_col))
            self._com_end_range(document).InsertParagraphAfter()
            return
        if kind == "image":
            image_path = self._document_resolve_path(block.get("path"), mode="read")
            shape = document.InlineShapes.AddPicture(FileName=image_path, LinkToFile=False, SaveWithDocument=True, Range=target)
            if block.get("width_cm") is not None:
                shape.Width = self._cm_to_points(block["width_cm"])
            if block.get("height_cm") is not None:
                shape.Height = self._cm_to_points(block["height_cm"])
            try:
                shape.AlternativeText = str(block.get("alt_text") or "")
                shape.Title = str(block.get("title") or "")
            except Exception:
                pass
            target.ParagraphFormat.Alignment = {"left": 0, "center": 1, "right": 2}.get(str(block.get("alignment") or "center").lower(), 1)
            target.InsertParagraphAfter()
            return
        if kind == "page_break":
            target.InsertBreak(7)
            return
        if kind == "section_break":
            target.InsertBreak({"continuous": 3, "even_page": 4, "odd_page": 5}.get(str(block.get("start") or "new_page").lower(), 2))
            if block.get("page"):
                try:
                    affected = [target.Sections(1)]
                except Exception:
                    affected = [document.Sections(document.Sections.Count)]
                self._com_set_page_layout(document, block.get("page") or {}, sections=affected)
            return
        if kind == "toc":
            document.TablesOfContents.Add(
                Range=target, UseHeadingStyles=True,
                UpperHeadingLevel=int(block.get("upper_level") or 1),
                LowerHeadingLevel=int(block.get("lower_level") or 3),
                UseHyperlinks=True, HidePageNumbersInWeb=True,
            )
            target.InsertParagraphAfter()
            return
        if kind in {"header", "footer"}:
            self._com_set_header_footer(document, kind, block)
            return
        if kind == "comment":
            comment_range = self._com_select_range(document, block.get("selector") or {})
            document.Comments.Add(Range=comment_range, Text=str(block.get("text") or ""))
            return
        if kind in {"footnote", "endnote"}:
            note_range = self._com_note_range(document, block)
            collection = document.Footnotes if kind == "footnote" else document.Endnotes
            kwargs = {"Range": note_range, "Text": str(block.get("text") or "")}
            if block.get("reference_text"):
                kwargs["Reference"] = str(block.get("reference_text"))
            collection.Add(**kwargs)
            return
        if kind == "content_control":
            control_type = _friendly_enum(
                block.get("control_type"), _WORD_CONTENT_CONTROL_TYPES, "control_type", 0
            )
            control = document.ContentControls.Add(Type=control_type, Range=target)
            control.Title = str(block.get("title") or "")
            control.Tag = str(block.get("tag") or "")
            if control_type == 8:
                control.Checked = bool(block.get("checked", False))
            else:
                for raw_item in block.get("items") or []:
                    item = raw_item if isinstance(raw_item, dict) else {"display": raw_item, "value": raw_item}
                    control.DropdownListEntries.Add(
                        Text=str(item.get("display") or item.get("text") or item.get("value") or ""),
                        Value=str(item.get("value") or item.get("display") or item.get("text") or ""),
                    )
                if control_type == 6 and block.get("date_format"):
                    control.DateDisplayFormat = str(block.get("date_format"))
                text = str(block.get("text") or block.get("placeholder") or "")
                if text:
                    control.Range.Text = text
            self._com_apply_font(control.Range.Font, block)
            self._com_apply_paragraph_format(
                control.Range.ParagraphFormat, block, range_obj=control.Range
            )
            control.Range.LanguageID = 1037
            return
        if kind == "text_box":
            shape = document.Shapes.AddTextbox(
                Orientation=1,
                Left=self._cm_to_points(block.get("left_cm", 2)),
                Top=self._cm_to_points(block.get("top_cm", 2)),
                Width=self._cm_to_points(block.get("width_cm", 8)),
                Height=self._cm_to_points(block.get("height_cm", 3)),
                Anchor=target,
            )
            shape.TextFrame.TextRange.Text = str(block.get("text") or "")
            self._com_apply_font(shape.TextFrame.TextRange.Font, block)
            self._com_apply_paragraph_format(
                shape.TextFrame.TextRange.ParagraphFormat, block,
                range_obj=shape.TextFrame.TextRange,
            )
            return
        if kind == "shape":
            shape = document.Shapes.AddShape(
                Type=int(block.get("shape_type") or 1),
                Left=self._cm_to_points(block.get("left_cm", 2)),
                Top=self._cm_to_points(block.get("top_cm", 2)),
                Width=self._cm_to_points(block.get("width_cm", 5)),
                Height=self._cm_to_points(block.get("height_cm", 3)),
                Anchor=target,
            )
            if block.get("text"):
                shape.TextFrame.TextRange.Text = str(block["text"])
                self._com_apply_font(shape.TextFrame.TextRange.Font, block)
                self._com_apply_paragraph_format(
                    shape.TextFrame.TextRange.ParagraphFormat, block,
                    range_obj=shape.TextFrame.TextRange,
                )
            if block.get("fill"):
                shape.Fill.ForeColor.RGB = self._com_rgb(block["fill"])
            if block.get("line_color"):
                shape.Line.ForeColor.RGB = self._com_rgb(block["line_color"])
            if block.get("line_weight_pt") is not None:
                shape.Line.Weight = _clamp(block["line_weight_pt"], 0, 20, 1)
            if block.get("rotation") is not None:
                shape.Rotation = _clamp(block["rotation"], -360, 360, 0)
            return
        if kind == "chart":
            chart_type = _friendly_enum(block.get("chart_type"), _WORD_CHART_TYPES, "chart_type", 51)
            inline_shape = document.InlineShapes.AddChart2(-1, chart_type, target)
            if block.get("width_cm") is not None:
                inline_shape.Width = self._cm_to_points(block.get("width_cm"))
            if block.get("height_cm") is not None:
                inline_shape.Height = self._cm_to_points(block.get("height_cm"))
            try:
                inline_shape.AlternativeText = str(block.get("alt_text") or "")
                inline_shape.Title = str(block.get("title") or "")
            except Exception:
                pass
            chart = inline_shape.Chart
            title = str(block.get("title") or "").strip()
            chart.HasTitle = bool(title)
            if title:
                chart.ChartTitle.Text = title
            series_specs = block.get("series") or []
            if series_specs:
                series_collection = chart.SeriesCollection()
                while int(series_collection.Count) > 0:
                    series_collection(1).Delete()
                categories = tuple(block.get("categories") or [])
                for index, series_spec in enumerate(series_specs, 1):
                    series = series_collection.NewSeries()
                    series.Name = str(series_spec.get("name") or f"Series {index}")
                    series.Values = tuple(series_spec.get("values") or [])
                    if categories:
                        series.XValues = categories
            show_legend = bool(block.get("show_legend", len(series_specs) > 1))
            chart.HasLegend = show_legend
            if show_legend and block.get("legend_position") is not None:
                chart.Legend.Position = _friendly_enum(
                    block.get("legend_position"), _WORD_LEGEND_POSITIONS, "legend_position", -4107
                )
            return
        if kind == "equation":
            text = str(block.get("text") or "")
            target.Text = text
            equation_range = document.OMaths.Add(target)
            equation_range.OMaths(1).BuildUp()
            return
        if kind == "advanced_com":
            self._com_advanced(document, block, args)
            return
        raise ValueError(f"Unsupported COM block type: {kind}")

    def _com_set_header_footer(self, document, kind, spec):
        spec = _normalized_format_spec(_merge_nested_spec(spec, "content"))
        for section_index in range(1, int(document.Sections.Count) + 1):
            section = document.Sections(section_index)
            if "different_first_page" in spec:
                section.PageSetup.DifferentFirstPageHeaderFooter = bool(spec["different_first_page"])
            stories = section.Headers if kind == "header" else section.Footers
            story_type = {"primary": 1, "first": 2, "even": 3}.get(str(spec.get("variant") or "primary").lower(), 1)
            story = stories(story_type)
            story.LinkToPrevious = bool(spec.get("link_to_previous", False))
            story.Range.Text = str(spec.get("text") or "")
            self._com_apply_font(story.Range.Font, spec)
            self._com_apply_paragraph_format(
                story.Range.ParagraphFormat, spec, range_obj=story.Range
            )
            if spec.get("page_number"):
                story.Range.InsertAfter(" ")
                insertion = story.Range.Duplicate
                insertion.Collapse(0)
                story.Range.Fields.Add(Range=insertion, Type=33, PreserveFormatting=True)

    @staticmethod
    def _com_subrange(base_range, relative_start, relative_end):
        result = base_range.Duplicate
        base_start = int(base_range.Start)
        result.SetRange(base_start + int(relative_start), base_start + int(relative_end))
        return result

    def _com_find_literal_range(self, base_range, text, *, occurrence=1, case_sensitive=False, whole_word=False):
        spans = _literal_match_spans(
            str(base_range.Text or ""),
            text,
            case_sensitive=case_sensitive,
            whole_word=whole_word,
        )
        occurrence = max(1, int(occurrence or 1))
        if len(spans) < occurrence:
            preview = str(text or "").replace("\r", " ").replace("\n", " ")[:160]
            raise ValueError(f"Text selector did not match occurrence {occurrence}: {preview}")
        start, end = spans[occurrence - 1]
        return self._com_subrange(base_range, start, end)

    def _com_replace_long_literal(
        self,
        target,
        old,
        new,
        *,
        case_sensitive=False,
        whole_word=False,
        replace_all=True,
    ):
        spans = _literal_match_spans(
            str(target.Text or ""),
            old,
            case_sensitive=case_sensitive,
            whole_word=whole_word,
        )
        if not spans:
            return 0
        selected = spans if replace_all else spans[:1]
        replacement = str(new or "").replace("\r\n", "\r").replace("\n", "\r")
        # Replace from the end so earlier absolute character positions remain
        # valid even when replacement lengths differ from their matches.
        for start, end in reversed(selected):
            cancellation_check = getattr(self, "_raise_if_cancelled", None)
            if callable(cancellation_check):
                cancellation_check()
            replacement_range = self._com_subrange(target, start, end)
            replacement_range.Text = replacement
        return len(selected)

    def _com_select_range(self, document, selector):
        selector = selector if isinstance(selector, dict) else {}
        if selector.get("bookmark"):
            name = str(selector["bookmark"])
            if not document.Bookmarks.Exists(name):
                raise ValueError(f"Bookmark not found: {name}")
            return document.Bookmarks(name).Range.Duplicate
        if selector.get("paragraph_index") is not None:
            index = int(selector["paragraph_index"]) + 1
            if index < 1 or index > int(document.Paragraphs.Count):
                raise IndexError("paragraph_index is out of range.")
            return document.Paragraphs(index).Range.Duplicate
        if selector.get("table_index") is not None:
            table = document.Tables(int(selector["table_index"]) + 1)
            if selector.get("row") is not None and selector.get("column") is not None:
                result = table.Cell(int(selector["row"]) + 1, int(selector["column"]) + 1).Range.Duplicate
                result.End = max(result.Start, result.End - 1)
                return result
            return table.Range.Duplicate
        if selector.get("start") is not None or selector.get("end") is not None:
            start = int(selector.get("start") or 0)
            end = int(selector.get("end") if selector.get("end") is not None else start)
            if start < 0 or end < start or end > int(document.Content.End):
                raise ValueError("Invalid character range selector.")
            return document.Range(start, end)
        find_text = str(selector.get("find") or selector.get("text") or "")
        if find_text:
            occurrence = max(1, int(selector.get("occurrence") or 1))
            if len(find_text) > _WORD_FIND_STRING_LIMIT:
                return self._com_find_literal_range(
                    document.Content.Duplicate,
                    find_text,
                    occurrence=occurrence,
                    case_sensitive=bool(selector.get("case_sensitive", False)),
                    whole_word=bool(selector.get("whole_word", False)),
                )
            search = document.Content.Duplicate
            for _index in range(occurrence):
                search.Find.ClearFormatting()
                found = search.Find.Execute(
                    FindText=find_text,
                    MatchCase=bool(selector.get("case_sensitive", False)),
                    MatchWholeWord=bool(selector.get("whole_word", False)),
                    Forward=True,
                    Wrap=0,
                )
                if not found:
                    raise ValueError(f"Text selector did not match occurrence {occurrence}: {find_text}")
                if _index + 1 < occurrence:
                    search.SetRange(search.End, document.Content.End)
            return search
        return document.Content.Duplicate

    def _com_note_range(self, document, spec):
        selector = spec.get("selector") if isinstance(spec, dict) else None
        if isinstance(selector, dict) and selector:
            target = self._com_select_range(document, selector)
        else:
            target = self._com_end_range(document).Duplicate
        target.Collapse(0)
        return target

    def _com_apply_operation(self, document, operation, args):
        op = str(operation.get("op") or operation.get("type") or "").lower()
        if op in {"append", "append_blocks"}:
            blocks = operation.get("blocks") or []
            self._com_add_blocks(document, blocks, args)
            return {"op": op, "count": len(blocks)}
        if op in {"insert", "insert_blocks"}:
            target = self._com_select_range(document, operation.get("selector") or {})
            target.Collapse(1 if str(operation.get("position") or "before").lower() == "before" else 0)
            blocks = operation.get("blocks") or []
            self._com_add_blocks(document, blocks, args, target)
            return {"op": op, "count": len(blocks)}
        if op == "replace_text":
            target = self._com_select_range(document, operation.get("selector") or {})
            old = str(operation.get("find") or operation.get("old") or "")
            if not old:
                raise ValueError("replace_text requires find.")
            new = str(operation.get("replace") if "replace" in operation else operation.get("new") or "")
            if len(old) > _WORD_FIND_STRING_LIMIT or len(new) > _WORD_FIND_STRING_LIMIT:
                count = self._com_replace_long_literal(
                    target,
                    old,
                    new,
                    case_sensitive=bool(operation.get("case_sensitive", False)),
                    whole_word=bool(operation.get("whole_word", False)),
                    replace_all=bool(operation.get("replace_all", True)),
                )
                return {"op": op, "matched": bool(count), "count": count, "method": "range"}
            replaced = target.Find.Execute(
                FindText=old, MatchCase=bool(operation.get("case_sensitive", False)),
                MatchWholeWord=bool(operation.get("whole_word", False)),
                ReplaceWith=new, Replace=2 if operation.get("replace_all", True) else 1,
                Forward=True, Wrap=0,
            )
            return {"op": op, "matched": bool(replaced), "method": "word_find"}
        if op in {"format", "format_range", "format_paragraph"}:
            target = self._com_select_range(document, operation.get("selector") or {})
            spec = _normalized_format_spec(operation)
            if spec.get("style"):
                target.Style = self._com_style_ref(spec["style"])
            self._com_apply_font(target.Font, spec)
            self._com_apply_paragraph_format(
                target.ParagraphFormat, spec, range_obj=target
            )
            target.LanguageID = 1037
            return {"op": op, "status": "ok"}
        if op in {"delete", "delete_range", "delete_paragraph"}:
            target = self._com_select_range(document, operation.get("selector") or operation)
            target.Delete()
            return {"op": op, "status": "ok"}
        if op == "set_page_layout":
            self._com_set_page_layout(document, operation)
            return {"op": op, "status": "ok"}
        if op == "define_style":
            self._com_define_style(document, operation)
            return {"op": op, "status": "ok"}
        if op in {"set_header", "set_footer"}:
            self._com_set_header_footer(document, op.split("_")[1], operation)
            return {"op": op, "status": "ok"}
        if op == "set_properties":
            self._com_set_properties(document, operation.get("properties") or operation)
            return {"op": op, "status": "ok"}
        if op == "update_fields":
            self._com_prepare_toc_styles(document)
            document.Fields.Update()
            for index in range(1, int(document.TablesOfContents.Count) + 1):
                document.TablesOfContents(index).Update()
            return {"op": op, "status": "ok"}
        if op == "track_changes":
            document.TrackRevisions = bool(operation.get("enabled", True))
            return {"op": op, "enabled": bool(document.TrackRevisions)}
        if op == "accept_all_changes":
            document.AcceptAllRevisions()
            return {"op": op, "status": "ok"}
        if op == "reject_all_changes":
            document.RejectAllRevisions()
            return {"op": op, "status": "ok"}
        if op == "protect":
            document.Protect(
                Type=int(operation.get("protection_type") or 3),
                NoReset=True,
                Password=str(operation.get("password") or ""),
                UseIRM=False,
                EnforceStyleLock=bool(operation.get("enforce_style_lock", False)),
            )
            return {"op": op, "status": "ok"}
        if op == "unprotect":
            document.Unprotect(Password=str(operation.get("password") or ""))
            return {"op": op, "status": "ok"}
        if op in {"add_comment", "add_footnote", "add_endnote"}:
            text = str(operation.get("text") or "")
            if op == "add_comment":
                target = self._com_select_range(document, operation.get("selector") or {})
                document.Comments.Add(Range=target, Text=text)
            elif op == "add_footnote":
                kwargs = {"Range": self._com_note_range(document, operation), "Text": text}
                if operation.get("reference_text"):
                    kwargs["Reference"] = str(operation.get("reference_text"))
                document.Footnotes.Add(**kwargs)
            else:
                kwargs = {"Range": self._com_note_range(document, operation), "Text": text}
                if operation.get("reference_text"):
                    kwargs["Reference"] = str(operation.get("reference_text"))
                document.Endnotes.Add(**kwargs)
            return {"op": op, "status": "ok"}
        if op in {"add_text_box", "add_shape", "add_chart", "add_equation"}:
            block_type = op.replace("add_", "")
            self._com_add_block(document, {**operation, "type": block_type}, args)
            return {"op": op, "status": "ok"}
        if op == "insert_file":
            file_path = self._document_resolve_path(operation.get("path"), mode="read")
            target = self._com_select_range(document, operation.get("selector") or {})
            target.Collapse(0)
            target.InsertFile(FileName=file_path, ConfirmConversions=False, Link=False, Attachment=False)
            return {"op": op, "status": "ok", "path": file_path}
        if op == "advanced_com":
            return {"op": op, "result": self._com_advanced(document, operation, args)}
        raise ValueError(f"Unsupported COM edit operation: {op}")

    def _com_advanced(self, document, operation, args):
        if not args.get("allow_advanced_com"):
            raise PermissionError("advanced_com requires allow_advanced_com=true in the document_manager call.")
        root_name = str(operation.get("root") or "document").lower()
        if root_name == "document":
            target = document
        elif root_name == "application":
            target = document.Application
        elif root_name == "options":
            target = document.Application.Options
        else:
            raise ValueError("advanced_com root must be document, application, or options.")
        for step in operation.get("path") or []:
            if isinstance(step, str):
                member, index = step, None
            elif isinstance(step, dict):
                member, index = str(step.get("member") or ""), step.get("index")
            else:
                raise ValueError("advanced_com path entries must be strings or {member,index} objects.")
            self._validate_com_member(member)
            target = getattr(target, member)
            if index is not None:
                target = target.Item(index) if hasattr(target, "Item") else target(index)
        member = str(operation.get("member") or "")
        self._validate_com_member(member)
        mode = str(operation.get("mode") or "get").lower()
        if mode == "get":
            result = getattr(target, member)
        elif mode == "set":
            setattr(target, member, operation.get("value"))
            result = getattr(target, member)
        elif mode == "call":
            method = getattr(target, member)
            values = operation.get("args") or []
            named = operation.get("kwargs") or {}
            if not isinstance(values, list) or not isinstance(named, dict):
                raise ValueError("advanced_com call args must be an array and kwargs an object.")
            result = method(*values, **named)
        else:
            raise ValueError("advanced_com mode must be get, set, or call.")
        return _safe_com_result(result)

    @staticmethod
    def _validate_com_member(member):
        name = str(member or "").strip()
        if not re.fullmatch(r"[A-Za-z][A-Za-z0-9_]{0,79}", name):
            raise ValueError(f"Unsafe COM member name: {name!r}")
        folded = name.casefold()
        blocked_tokens = (
            "vbproject", "macro", "oleobject", "commandbar", "addins", "wordbasic",
            "dde", "organizer", "print", "sendmail", "sendfax", "followhyperlink",
            "filedialog",
        )
        if folded in _ADVANCED_COM_BLOCKED or any(token in folded for token in blocked_tokens):
            raise PermissionError(f"COM member '{name}' is blocked by Smarti document safety policy.")

    def _com_save_document(self, document, temp_path, desired_path):
        fmt = os.path.splitext(desired_path)[1].lower().lstrip(".") or "docx"
        if fmt in _WORD_FIXED_FORMATS:
            document.ExportAsFixedFormat(OutputFileName=temp_path, ExportFormat=_WORD_FIXED_FORMATS[fmt], OpenAfterExport=False)
        elif fmt in _WORD_FORMATS:
            document.SaveAs2(
                FileName=temp_path,
                FileFormat=_WORD_FORMATS[fmt],
                AddToRecentFiles=False,
                EmbedTrueTypeFonts=True,
                AddBiDiMarks=True,
            )
        else:
            raise ValueError(f"Unsupported Word output format: {fmt}")

    # --------------------- inspect / export / render ----------------------

    def _document_inspect(self, args, engine):
        path = self._document_resolve_path(args.get("path"), mode="read")
        if os.path.splitext(path)[1].lower() == ".pdf":
            return self._inspect_pdf(path, args)
        if engine == "com":
            return self._com_inspect_document(path, args)
        if os.path.splitext(path)[1].lower() != ".docx":
            raise ValueError("The independent python inspect engine supports DOCX only; use engine='com' for other Word formats.")
        return self._python_inspect_document(path, args)

    def _com_inspect_document(self, path, args):
        paragraph_limit = max(0, min(1000, int(args.get("paragraph_limit") or 200)))
        include_text = bool(args.get("include_text", True))
        with _WordComSession(args.get("timeout_seconds", 180), args.get("visible", False)) as session:
            document = session.open_document(path, read_only=True, password=str(args.get("password") or ""))
            paragraph_count = int(document.Paragraphs.Count)
            paragraphs = []
            if include_text:
                for index in range(1, min(paragraph_count, paragraph_limit) + 1):
                    paragraph = document.Paragraphs(index)
                    text = str(paragraph.Range.Text or "").rstrip("\r\x07")
                    paragraphs.append({
                        "index": index - 1,
                        "style": str(paragraph.Style),
                        "text": text[:4000],
                        "alignment": int(paragraph.Format.Alignment),
                        "reading_order": "rtl" if int(paragraph.Format.ReadingOrder) == 0 else "ltr",
                    })
            sections = []
            for index in range(1, int(document.Sections.Count) + 1):
                setup = document.Sections(index).PageSetup
                sections.append({
                    "width_pt": round(float(setup.PageWidth), 3),
                    "height_pt": round(float(setup.PageHeight), 3),
                    "orientation": "landscape" if int(setup.Orientation) == 1 else "portrait",
                    "direction": "rtl" if int(setup.SectionDirection) == 0 else "ltr",
                    "margins_pt": {
                        "top": round(float(setup.TopMargin), 3),
                        "bottom": round(float(setup.BottomMargin), 3),
                        "left": round(float(setup.LeftMargin), 3),
                        "right": round(float(setup.RightMargin), 3),
                    },
                })
            properties = {}
            for key, name in (
                ("title", "Title"), ("subject", "Subject"), ("author", "Author"),
                ("keywords", "Keywords"), ("comments", "Comments"), ("category", "Category"),
            ):
                try:
                    properties[key] = str(document.BuiltInDocumentProperties(name).Value or "")
                except Exception:
                    pass
            rtl_sample_limit = min(paragraph_count, 2000)
            rtl_count = sum(
                1 for index in range(1, rtl_sample_limit + 1)
                if int(document.Paragraphs(index).Format.ReadingOrder) == 0
            )
            result = {
                "status": "ok",
                "engine": "com",
                "path": path,
                "format": os.path.splitext(path)[1].lower().lstrip("."),
                "paragraph_count": paragraph_count,
                "table_count": int(document.Tables.Count),
                "section_count": int(document.Sections.Count),
                "inline_shape_count": int(document.InlineShapes.Count),
                "shape_count": int(document.Shapes.Count),
                "field_count": int(document.Fields.Count),
                "comment_count": int(document.Comments.Count),
                "footnote_count": int(document.Footnotes.Count),
                "endnote_count": int(document.Endnotes.Count),
                "content_control_count": int(document.ContentControls.Count),
                "rtl_paragraph_count": rtl_count,
                "rtl_count_sampled_paragraphs": rtl_sample_limit,
                "track_revisions": bool(document.TrackRevisions),
                "protection_type": int(document.ProtectionType),
                "properties": properties,
                "sections": sections,
                "paragraphs": paragraphs,
                "truncated": paragraph_count > paragraph_limit,
            }
            session.close_document(False)
            return result

    def _python_inspect_document(self, path, args):
        api = self._python_imports()
        document = api["docx"].Document(path)
        paragraph_limit = max(0, min(1000, int(args.get("paragraph_limit") or 200)))
        include_text = bool(args.get("include_text", True))
        paragraphs = []
        if include_text:
            for index, paragraph in enumerate(document.paragraphs[:paragraph_limit]):
                paragraphs.append({
                    "index": index,
                    "style": paragraph.style.name if paragraph.style else "",
                    "text": paragraph.text[:4000],
                })
        sections = []
        for section in document.sections:
            sections.append({
                "width_cm": round(section.page_width.cm, 3) if section.page_width else None,
                "height_cm": round(section.page_height.cm, 3) if section.page_height else None,
                "orientation": str(section.orientation),
                "margins_cm": {
                    "top": round(section.top_margin.cm, 3) if section.top_margin else None,
                    "bottom": round(section.bottom_margin.cm, 3) if section.bottom_margin else None,
                    "left": round(section.left_margin.cm, 3) if section.left_margin else None,
                    "right": round(section.right_margin.cm, 3) if section.right_margin else None,
                },
            })
        with zipfile.ZipFile(path, "r") as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8", "replace")
            names = archive.namelist()
        return {
            "status": "ok",
            "path": path,
            "format": os.path.splitext(path)[1].lower().lstrip("."),
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "section_count": len(document.sections),
            "inline_shape_count": len(document.inline_shapes),
            "comment_part": "word/comments.xml" in names,
            "footnote_part": "word/footnotes.xml" in names,
            "endnote_part": "word/endnotes.xml" in names,
            "field_count": document_xml.count("w:fldChar"),
            "rtl_paragraph_count": document_xml.count("w:bidi"),
            "sections": sections,
            "paragraphs": paragraphs,
            "truncated": len(document.paragraphs) > paragraph_limit,
        }

    @staticmethod
    def _inspect_pdf(path, args):
        fitz = _pymupdf_module()
        document = fitz.open(path)
        try:
            return {
                "status": "ok", "path": path, "format": "pdf", "page_count": document.page_count,
                "pages": [
                    {"page": index + 1, "width_pt": page.rect.width, "height_pt": page.rect.height, "text_chars": len(page.get_text())}
                    for index, page in enumerate(document)
                ],
            }
        finally:
            document.close()

    def _document_export(self, args, engine):
        source = self._document_resolve_path(args.get("path"), mode="read")
        fmt = str(args.get("format") or os.path.splitext(str(args.get("output_path") or ""))[1] or "pdf").lower().lstrip(".")
        output = self._document_resolve_path(
            args.get("output_path"), mode="write",
            default_name=os.path.splitext(os.path.basename(source))[0] + "." + fmt,
        )
        self._document_prepare_output(output, bool(args.get("overwrite")), backup_existing=True)
        return self._document_export_internal(source, output, fmt, engine, args)

    def _document_export_internal(self, source, output, fmt, engine, args):
        if os.path.normcase(source) == os.path.normcase(output):
            raise ValueError("Export output_path must differ from the source path.")
        temp = self._document_temp_path(output)
        try:
            if engine == "com":
                with _WordComSession(args.get("timeout_seconds", 180), args.get("visible", False)) as session:
                    document = session.open_document(source, read_only=True, password=str(args.get("password") or ""))
                    try:
                        self._com_prepare_toc_styles(document)
                        document.Fields.Update()
                    except Exception:
                        pass
                    self._com_save_document(document, temp, output)
                    session.close_document(False)
            elif engine == "libreoffice":
                produced = self._libreoffice_convert(source, os.path.dirname(temp), fmt)
                shutil.move(produced, temp)
            elif engine == "python":
                if fmt == "docx":
                    shutil.copy2(source, temp)
                elif fmt == "txt":
                    api = self._python_imports()
                    document = api["docx"].Document(source)
                    with open(temp, "w", encoding="utf-8-sig", newline="") as handle:
                        handle.write("\n".join(paragraph.text for paragraph in document.paragraphs))
                else:
                    raise ValueError(f"python engine cannot export {fmt}; use com or libreoffice.")
            else:
                raise ValueError(f"Unsupported export engine: {engine}")
            os.replace(temp, output)
        finally:
            if os.path.exists(temp):
                os.remove(temp)
        return {"status": "exported", "engine": engine, "format": fmt, "output_path": output}

    def _libreoffice_convert(self, source, output_dir, fmt):
        soffice = self._find_soffice()
        if not soffice:
            raise RuntimeError("LibreOffice/soffice is unavailable.")
        os.makedirs(output_dir, exist_ok=True)
        profile = os.path.join(tempfile.gettempdir(), f"smarti-lo-{uuid.uuid4().hex}")
        profile_url = pathlib.Path(profile).as_uri()
        command = [
            soffice, "--headless", "--nologo", "--nodefault", "--nofirststartwizard",
            f"-env:UserInstallation={profile_url}", "--convert-to", fmt, "--outdir", output_dir, source,
        ]
        try:
            completed = subprocess.run(command, capture_output=True, text=True, timeout=180, creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0))
            if completed.returncode != 0:
                raise RuntimeError((completed.stderr or completed.stdout or "LibreOffice conversion failed")[:2000])
            expected = os.path.join(output_dir, os.path.splitext(os.path.basename(source))[0] + "." + fmt)
            if not os.path.isfile(expected):
                candidates = sorted(pathlib.Path(output_dir).glob(f"*.{fmt}"), key=lambda item: item.stat().st_mtime, reverse=True)
                if not candidates:
                    raise RuntimeError("LibreOffice did not produce the requested output file.")
                expected = str(candidates[0])
            return expected
        finally:
            shutil.rmtree(profile, ignore_errors=True)

    def _document_render(self, args, engine):
        source = self._document_resolve_path(args.get("path"), mode="read")
        fitz = _pymupdf_module()
        stem = os.path.splitext(os.path.basename(source))[0]
        output_dir = self._document_resolve_path(
            args.get("output_dir"), mode="write",
            default_name=f"{stem}-render-{datetime.now().strftime('%Y%m%d-%H%M%S')}",
        )
        os.makedirs(output_dir, exist_ok=True)
        if os.path.splitext(source)[1].lower() == ".pdf":
            pdf_path = source
            export_info = None
        else:
            pdf_path = os.path.join(output_dir, stem + ".pdf")
            render_engine = engine
            if render_engine == "python":
                raise ValueError("The python engine can rasterize an existing PDF but cannot render DOCX directly. Use engine='auto', 'com', or 'libreoffice'.")
            export_info = self._document_export_internal(source, pdf_path, "pdf", render_engine, {**args, "overwrite": True})
        dpi = int(_clamp(args.get("dpi"), 72, 300, 144))
        page_limit = int(_clamp(args.get("page_limit"), 1, 500, 100))
        document = fitz.open(pdf_path)
        images = []
        warnings = []
        try:
            if document.page_count > page_limit:
                raise ValueError(f"Document has {document.page_count} pages; page_limit is {page_limit}.")
            matrix = fitz.Matrix(dpi / 72.0, dpi / 72.0)
            for index, page in enumerate(document):
                image_path = os.path.join(output_dir, f"page-{index + 1:03d}.png")
                pixmap = page.get_pixmap(matrix=matrix, alpha=False)
                pixmap.save(image_path)
                text_chars = len(page.get_text().strip())
                if text_chars == 0:
                    warnings.append(f"Page {index + 1} contains no extractable text; verify that it is intentional or image-only.")
                images.append({
                    "page": index + 1, "path": image_path,
                    "width_px": pixmap.width, "height_px": pixmap.height,
                    "text_chars": text_chars,
                })
        finally:
            document.close()
        pdf_retained = True
        if export_info is not None and not bool(args.get("include_pdf", True)):
            os.remove(pdf_path)
            pdf_retained = False
            export_info = {**export_info, "output_path": None, "retained": False}
        elif export_info is not None:
            export_info = {**export_info, "retained": True}
        return {
            "status": "rendered",
            "engine": engine,
            "source_path": source,
            "pdf_path": pdf_path if pdf_retained else None,
            "pdf_retained": pdf_retained,
            "output_dir": output_dir,
            "page_count": len(images),
            "pages": images,
            "warnings": warnings,
            "visual_qa_required": True,
            "next_step": "Use document_manager action=visual_qa on every page PNG, then edit and re-render if any visual defect is found.",
            "export": export_info,
        }

    def _document_compare(self, args):
        original_path = self._document_resolve_path(args.get("path"), mode="read")
        revised_path = self._document_resolve_path(args.get("other_path"), mode="read")
        output = self._document_resolve_path(
            args.get("output_path"), mode="write",
            default_name=f"comparison-{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx",
        )
        self._document_prepare_output(output, bool(args.get("overwrite")), backup_existing=True)
        temp = self._document_temp_path(output)
        timeout = args.get("timeout_seconds", 240)
        with _WordComSession(timeout, args.get("visible", False)) as session:
            original = session.app.Documents.Open(FileName=original_path, ReadOnly=True, AddToRecentFiles=False, Visible=False)
            revised = session.app.Documents.Open(FileName=revised_path, ReadOnly=True, AddToRecentFiles=False, Visible=False)
            compared = None
            try:
                compared = session.app.CompareDocuments(
                    OriginalDocument=original,
                    RevisedDocument=revised,
                    Destination=2,
                    Granularity=1,
                    CompareFormatting=True,
                    CompareCaseChanges=True,
                    CompareWhitespace=True,
                    CompareTables=True,
                    CompareHeaders=True,
                    CompareFootnotes=True,
                    CompareTextboxes=True,
                    CompareFields=True,
                    CompareComments=True,
                    CompareMoves=True,
                    RevisedAuthor=str(args.get("revised_author") or "Smarti"),
                    IgnoreAllComparisonWarnings=True,
                )
                compared.SaveAs2(FileName=temp, FileFormat=16, AddToRecentFiles=False, EmbedTrueTypeFonts=True, AddBiDiMarks=True)
            finally:
                if compared is not None:
                    try:
                        compared.Close(SaveChanges=0)
                    except Exception:
                        pass
                try:
                    revised.Close(SaveChanges=0)
                except Exception:
                    pass
                try:
                    original.Close(SaveChanges=0)
                except Exception:
                    pass
        os.replace(temp, output)
        return {"status": "compared", "engine": "com", "output_path": output, "original_path": original_path, "revised_path": revised_path}
